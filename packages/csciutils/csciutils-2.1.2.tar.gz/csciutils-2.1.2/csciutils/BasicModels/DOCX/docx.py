from docx import Document
import re

class DOCX():
    def __init__(self) -> None:
        pass

    @staticmethod
    def find_target_in_string(A, target:None)->list:
        if target is None:
            return []
        # 使用正则表达式查找所有匹配的 target
        matches = [(match.start(), match.end()) for match in re.finditer(re.escape(target), A)]
        return matches
    
    @staticmethod
    def find_overlapping_ranges(run_start_index:int, run_end_index:int, target_indexes:list)->list:
        overlaps = []
        if not target_indexes:
            pass
        else:
            for target_index_tuple in target_indexes:
                # 找到重叠的开始和结束索引
                start_overlap = max(run_start_index, target_index_tuple[0])
                end_overlap = min(run_end_index, target_index_tuple[1])
                
                # 如果开始索引小于或等于结束索引，说明有重叠部分
                if start_overlap < end_overlap:
                    overlaps.append({
                        "overlap": (start_overlap, end_overlap),
                        "start": start_overlap == target_index_tuple[0]
                    })
        return overlaps

    @staticmethod
    def replace_in_paragraph(paragraph, target, replacement = None):
        """
        - replacement: str | default = None | 如果是None就不做任何变动，直接跳过当前函数；如果要删除原来的信息，replacement为""
        """
        if replacement is None:
            return paragraph
        elif target in paragraph.text:
            print(f'found target = {target} in {paragraph.text}')
            target_indexes = DOCX.find_target_in_string(paragraph.text, target)
            # print(target_indexes)
            run_start_index = 0
            run_end_index = 0
            for run in paragraph.runs:
                run_start_index = run_end_index
                run_end_index = run_start_index + len(run.text)
                # print(f"run_text = {run.text}, start_index = {run_start_index}, end_index = {run_end_index}")
                overlaps = DOCX.find_overlapping_ranges(run_start_index=run_start_index, run_end_index=run_end_index, target_indexes=target_indexes)
                # print(overlaps)
                if overlaps:
                    for overlap in overlaps:
                        overlap_start, overlap_end = overlap.get('overlap')
                        if overlap.get('start'):
                            # print(f"Replacing {run.text[overlap_start - run_start_index: overlap_end-run_start_index]} with {replacement}")
                            run.text = run.text[0:overlap_start - run_start_index] + replacement + run.text[overlap_end-run_start_index:]
                        else:
                            # print(f"Replacing {run.text[overlap_start - run_start_index: overlap_end-run_start_index]} with {''}")
                            run.text = run.text[0:overlap_start - run_start_index] + '' + run.text[overlap_end-run_start_index:]
        return paragraph

    @staticmethod
    def replace_in_doc(doc, target, replacement):
        print(f"Searching {target} in Text.")
        for paragraph in doc.paragraphs:
            paragraph = DOCX.replace_in_paragraph(paragraph, target, replacement)
        print(f"Searching {target} in Tables.")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph = DOCX.replace_in_paragraph(paragraph, target, replacement)
        print(f"Searching {target} in Footers.")
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph = DOCX.replace_in_paragraph(paragraph, target, replacement)
        return doc