import os
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def _extract_text_from_run(run: Run) -> str:
    """
    Extract text from a docx Run object with basic formatting.

    Args:
        run: docx Run object

    Returns:
        Formatted text string
    """
    text = run.text
    if not text:
        return ""

    # Apply basic formatting
    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"_{text}_"

    return text


def _extract_text_from_paragraph(paragraph: Paragraph) -> str:
    """
    Extract text from a docx Paragraph object.

    Args:
        paragraph: docx Paragraph object

    Returns:
        Formatted text string
    """
    if not paragraph.runs:
        return paragraph.text

    text_parts = []
    for run in paragraph.runs:
        text_parts.append(_extract_text_from_run(run))

    text = "".join(text_parts)

    # Handle paragraph alignment (basic)
    if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        text = f"<center>{text}</center>"
    elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        text = f"<div align='right'>{text}</div>"

    return text


def _extract_table_to_markdown(table: Table) -> str:
    """
    Convert a docx Table to markdown format.

    Args:
        table: docx Table object

    Returns:
        Markdown formatted table string
    """
    if not table.rows:
        return ""

    markdown_rows = []

    # Process each row
    for row_idx, row in enumerate(table.rows):
        row_cells = []

        for cell in row.cells:
            # Extract text from all paragraphs in the cell
            cell_text_parts = []
            for paragraph in cell.paragraphs:
                para_text = _extract_text_from_paragraph(paragraph)
                if para_text.strip():
                    cell_text_parts.append(para_text.strip())

            # Join multiple paragraphs with line breaks
            cell_text = " ".join(cell_text_parts)
            # Clean up extra spaces and newlines
            cell_text = re.sub(r"\s+", " ", cell_text).strip()

            # Escape pipe characters in cell content
            cell_text = cell_text.replace("|", "\\|")

            row_cells.append(cell_text if cell_text else " ")

        # Create markdown table row
        markdown_rows.append("| " + " | ".join(row_cells) + " |")

        # Add header separator after first row
        if row_idx == 0:
            separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
            markdown_rows.append(separator)

    return "\n".join(markdown_rows)


# TODO: Fix lint issues in this function
def _process_document_element(element: Any) -> str:  # noqa: C901
    """
    Process a document element (paragraph or table) and return markdown.

    Args:
        element: Document element (Paragraph or Table)

    Returns:
        Markdown formatted string
    """
    if isinstance(element, Paragraph):
        text = _extract_text_from_paragraph(element)
        if not text.strip():
            return ""

        # Handle different paragraph styles
        style_name = element.style.name.lower()

        if "heading" in style_name:
            # Extract heading level
            level = 1
            if "heading 1" in style_name:
                level = 1
            elif "heading 2" in style_name:
                level = 2
            elif "heading 3" in style_name:
                level = 3
            elif "heading 4" in style_name:
                level = 4
            elif "heading 5" in style_name:
                level = 5
            elif "heading 6" in style_name:
                level = 6

            return f"{'#' * level} {text}"

        return text

    elif isinstance(element, Table):
        return _extract_table_to_markdown(element)

    return ""


def docx_to_markdown_txt(docx_path: str) -> str:  # noqa: C901
    """
    Convert a DOCX file to markdown format.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Markdown formatted string

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        Exception: For other processing errors
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        doc = Document(docx_path)
        markdown_content = []

        # Process the document
        for element in doc.element.body:
            if element.tag.endswith("p"):
                # It's a paragraph
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        markdown_text = _process_document_element(paragraph)
                        if markdown_text:
                            markdown_content.append(markdown_text)
                        break
            elif element.tag.endswith("tbl"):
                # It's a table
                for table in doc.tables:
                    if table._element == element:
                        markdown_text = _process_document_element(table)
                        if markdown_text:
                            markdown_content.append(markdown_text)
                        break

        # Join with double newlines for proper markdown formatting
        return "\n\n".join(markdown_content)

    except Exception as e:
        raise Exception(f"Error extracting key elements from DOCX file: {e!s}") from e
