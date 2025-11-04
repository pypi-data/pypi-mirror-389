
from typing import List, Dict, Any
import os
import markdown
from docx import Document

def write_markdown(title: str, sections: List[Dict[str, str]], out_dir: str) -> str:
    os.makedirs(out_dir, exist_ok=True)
    md_path = os.path.join(out_dir, "iso_validation_report.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        for sec in sections:
            f.write(f"## {sec.get('heading','')}\n\n")
            f.write(sec.get("body_md", ""))
            f.write("\n\n")
    return md_path

def export_docx(md_path: str, out_dir: str) -> str:
    # super-simple markdown→docx (headings + paragraphs)
    doc = Document()
    with open(md_path, "r", encoding="utf-8") as f:
        for line in f:
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            else:
                doc.add_paragraph(line.rstrip())
    docx_path = os.path.join(os.path.dirname(md_path), "iso_validation_report.docx")
    doc.save(docx_path)
    return docx_path
