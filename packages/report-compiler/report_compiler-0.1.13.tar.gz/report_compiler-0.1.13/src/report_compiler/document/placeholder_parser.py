"""
Placeholder detection and parsing for DOCX documents.
"""

import re
from typing import Dict, List, Any, Optional
from docx import Document
from ..core.config import Config
from ..utils.logging_config import get_module_logger


class PlaceholderParser:
    """Handles detection and parsing of PDF placeholders in DOCX documents."""
    
    def __init__(self):
        self.overlay_regex = Config.OVERLAY_REGEX
        self.insert_regex = Config.INSERT_REGEX
        self.image_regex = Config.IMAGE_REGEX
        self.logger = get_module_logger(__name__)
        
        # Cache for document parsing
        self._doc = None
        self._doc_path = None
    
    def find_all_placeholders(self, docx_path: str) -> Dict[str, List[Dict]]:
        """
        Find all placeholders in a DOCX document.
        
        Args:
            docx_path: Path to DOCX document
            
        Returns:
            Dict with 'table' and 'paragraph' placeholder lists
        """
        self._load_document(docx_path)
        
        table_placeholders = self._find_table_placeholders()
        paragraph_placeholders = self._find_paragraph_placeholders()
        
        return {
            'table': table_placeholders,
            'paragraph': paragraph_placeholders,
            'total': len(table_placeholders) + len(paragraph_placeholders)
        }
    
    def _load_document(self, docx_path: str) -> None:
        """Load document if not already loaded or path changed."""
        if self._doc is None or self._doc_path != docx_path:
            self._doc = Document(docx_path)
            self._doc_path = docx_path
    
    def _find_table_placeholders(self) -> List[Dict]:
        """
        Find PDF placeholders inside single-cell tables (overlay type).
        
        Returns:
            List of dictionaries containing table placeholder info
        """
        placeholders = []
        
        try:
            for table_idx, table in enumerate(self._doc.tables):
                
                # Only consider single-cell tables for overlay inserts
                if len(table._cells) == 1:
                    cell = table.cell(0, 0)  # Single-cell table has only one cell
                    cell_text = cell.text.strip()
                    
                    # Check if this cell contains an OVERLAY placeholder
                    overlay_match = self.overlay_regex.search(cell_text)
                    image_match = self.image_regex.search(cell_text)
                    
                    if overlay_match:
                        path_raw = overlay_match.group(1).strip()
                        params_string = overlay_match.group(2)
                        
                        # Parse parameters
                        params = self._parse_overlay_parameters(params_string)
                        
                        self.logger.info("Found table (overlay) placeholder for: %s", path_raw)
                        self.logger.debug("      • Page specification: page=%s", params['page'])
                        self.logger.debug("      • Content cropping: %s", 'enabled' if params['crop'] else 'disabled')
                        self.logger.debug("      • Table index: %d", table_idx)
                        
                        table_info = {
                            'type': 'table',
                            'subtype': 'overlay',
                            'file_path': path_raw,
                            'page_spec': params['page'],
                            'crop_enabled': params['crop'],
                            'table_index': table_idx,
                            'table_text': cell_text,
                            'source': f'table_{table_idx}',
                            'is_recursive_docx': False, # Overlays cannot be recursive
                        }
                        
                        placeholders.append(table_info)
                    
                    elif image_match:
                        path_raw = image_match.group(1).strip()
                        params_string = image_match.group(2)
                        
                        # Parse parameters (similar to overlay but no page spec)
                        params = self._parse_image_parameters(params_string)
                        
                        self.logger.info("Found table (image) placeholder for: %s", path_raw)
                        self.logger.debug("      • Width: %s", params.get('width', 'auto'))
                        self.logger.debug("      • Height: %s", params.get('height', 'auto'))
                        self.logger.debug("      • Table index: %d", table_idx)
                        
                        table_info = {
                            'type': 'table',
                            'subtype': 'image',
                            'file_path': path_raw,
                            'width': params.get('width'),
                            'height': params.get('height'),
                            'table_index': table_idx,
                            'table_text': cell_text,
                            'source': f'table_{table_idx}',
                            'is_recursive_docx': False, # Images cannot be recursive
                        }
                        
                        placeholders.append(table_info)
                
                else:                    # Multi-cell tables: scan but don't classify as overlay
                    has_insert = False
                    for row in table.rows:
                        for cell in row.cells:
                            if (self.overlay_regex.search(cell.text) or 
                                self.insert_regex.search(cell.text) or
                                self.image_regex.search(cell.text)):
                                has_insert = True
                                break
                        if has_insert:
                            break
                    
                    if has_insert:
                        rows = len(table.rows)
                        cols = len(table.columns)
                        self.logger.warning("Multi-cell table #%d (%dx%d) contains a placeholder but is skipped (not a valid overlay type).", table_idx, rows, cols)
        
        except Exception as e:
            self.logger.error("Error scanning for table placeholders: %s", e, exc_info=True)
        
        return placeholders
    
    def _find_paragraph_placeholders(self) -> List[Dict]:
        """
        Find PDF placeholders in regular paragraphs (merge type).
        
        Returns:
            List of dictionaries containing paragraph placeholder info
        """
        placeholders = []
        
        try:
            for para_idx, paragraph in enumerate(self._doc.paragraphs):
                para_text = paragraph.text.strip()
                
                # Look for INSERT placeholders (merge type)
                match = self.insert_regex.search(para_text)
                if match:
                    file_path_raw = match.group(1).strip()
                    page_spec = match.group(2)  # Optional page specification
                    
                    is_docx = file_path_raw.lower().endswith('.docx')

                    if is_docx:
                        self.logger.info("Found recursive DOCX insert placeholder for: %s", file_path_raw)
                        if page_spec:
                            self.logger.warning("      • Page specification '%s' is ignored for DOCX inserts.", page_spec)
                            page_spec = None
                    else:
                        self.logger.info("Found paragraph (merge) placeholder for: %s", file_path_raw)
                        if page_spec:
                            self.logger.debug("      • Page specification: %s", page_spec)

                    self.logger.debug("      • Paragraph index: %d", para_idx)
                    
                    placeholder_info = {
                        'type': 'paragraph',
                        'file_path': file_path_raw,
                        'page_spec': page_spec,
                        'is_recursive_docx': is_docx,
                        'paragraph_index': para_idx,
                        'paragraph_text': para_text,
                        'source': f'paragraph_{para_idx}',
                    }
                    
                    placeholders.append(placeholder_info)
        
        except Exception as e:
            self.logger.error("Error scanning for paragraph placeholders: %s", e, exc_info=True)
        
        return placeholders
    
    def _parse_overlay_parameters(self, params_string: Optional[str]) -> Dict[str, Any]:
        """
        Parse overlay parameters from the placeholder text.
        
        Args:
            params_string: Parameter string after the path (can be None)
            
        Returns:
            Dict with parsed parameters
        """
        result = {
            'page': None,
            'crop': Config.DEFAULT_CROP_ENABLED
        }
        
        if not params_string:
            return result
        
        # Split parameters by comma
        params = [p.strip() for p in params_string.split(',')]
        
        for param in params:
            if '=' in param:
                key, value = param.split('=', 1)
                key = key.strip().lower()
                value = value.strip()
                
                if key == 'page':
                    result['page'] = value
                elif key == 'crop':
                    result['crop'] = value.lower() in ('true', '1', 'yes', 'on', 'enabled')
            else:
                # Assume it's a page specification if no key
                if param and not result['page']:
                    result['page'] = param
        
        return result

    def _parse_image_parameters(self, params_string: Optional[str]) -> Dict[str, Any]:
        """
        Parse image parameters from the placeholder text.
        
        Args:
            params_string: Parameter string after the path (can be None)
            
        Returns:
            Dict with parsed parameters (width, height)
        """
        result = {}
        
        if not params_string:
            return result
        
        # Split parameters by comma
        params = [p.strip() for p in params_string.split(',')]
        
        for param in params:
            if '=' in param:
                key, value = param.split('=', 1)
                key = key.strip().lower()
                value = value.strip()
                
                if key == 'width':
                    result['width'] = value
                elif key == 'height':
                    result['height'] = value
        
        return result
