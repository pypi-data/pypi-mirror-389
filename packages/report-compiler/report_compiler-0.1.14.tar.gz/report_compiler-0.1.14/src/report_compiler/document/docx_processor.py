"""
DOCX document processing and modification utilities.
"""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from PIL import Image
from ..core.config import Config
from ..utils.logging_config import get_docx_logger
from ..utils.conversions import points_to_inches, emu_to_points
from ..utils import docx_emf_patch  # Ensure EMF support is patched in


class DocxProcessor:
    """Handles DOCX document modification and marker insertion."""

    def __init__(self):
        self.logger = get_docx_logger()

    def create_modified_docx(
        self, input_path: str, placeholders: Dict[str, List[Dict]], output_path: str
    ) -> Optional[Dict[int, Dict[str, float]]]:
        """
        Create a modified DOCX, replacing placeholders with markers and extracting table dimensions.

        Args:
            input_path: Path to the source DOCX file.
            placeholders: Dictionary of found placeholders.
            output_path: Path to save the modified DOCX.

        Returns:
            A dictionary mapping table indices to their dimensions, or None on failure.
        """
        try:
            self.logger.debug("  > Loading source DOCX: %s", os.path.basename(input_path))
            doc = Document(input_path)
            table_metadata = {}

            # Process paragraph placeholders first
            if placeholders.get('paragraph'):
                self._process_paragraph_placeholders(doc, placeholders['paragraph'])

            # Process table placeholders and get their dimensions
            if placeholders.get('table'):
                table_metadata = self._process_table_placeholders(doc, placeholders['table'])

            self.logger.debug("  > Saving modified DOCX to: %s", os.path.basename(output_path))
            doc.save(output_path)
            return table_metadata

        except Exception as e:
            self.logger.error("❌ Failed to create modified document: %s", e, exc_info=True)
            return None

    def _process_paragraph_placeholders(self, doc: Document, para_placeholders: List[Dict]):
        """Replace paragraph placeholders with merge markers."""
        self.logger.debug("  > Processing %d paragraph (merge) placeholders...", len(para_placeholders))
        for placeholder in para_placeholders:
            para_idx = placeholder['paragraph_index']
            marker = Config.get_merge_marker(placeholder['paragraph_index'])
            self.logger.debug("    - Replacing paragraph %d with marker: %s", para_idx, marker)
            if para_idx < len(doc.paragraphs):
                p = doc.paragraphs[para_idx]
                p.clear()
                # It's better to add a page break after the paragraph to ensure separation
                p.add_run(marker)
                p.add_run().add_break()
            else:
                self.logger.warning("    - Paragraph index %d is out of bounds.", para_idx)

    def _replicate_table_rows_for_overlay(self, table, num_pages: int, table_idx: int):
        """Replicate table rows for multi-page PDF overlays."""
        if num_pages <= 1:
            return

        self.logger.debug("      - Replicating table rows for %d pages.", num_pages)
        try:
            first_row = table.rows[0]
            row_height = first_row.height
            first_cell = first_row.cells[0]
            # In python-docx, cell width is not a direct property. It's inherited from the table's column definition.
            # We assume all columns in our 1x1 table are the same, so we don't need to set width explicitly on the new cell.

            for i in range(1, num_pages):  # Loop for additional pages
                page_num = i + 1
                new_row = table.add_row()
                if row_height:
                    new_row.height = row_height
                
                new_cell = new_row.cells[0]
                # new_cell.text = ''
                marker = Config.get_overlay_marker(table_idx, page_num)
                # p = new_cell.add_paragraph(marker)
                new_cell.paragraphs[0].clear()
                new_cell.paragraphs[0].add_run(marker)
                self.logger.debug("        - Added marker for page %d: %s", page_num, marker)

        except Exception as e:
            self.logger.error("      - Failed to replicate table rows: %s", e, exc_info=True)

    def _get_table_dimensions_in_points(self, table, doc: Document) -> Dict[str, float]:
        """
        Calculates the width and height of a table in points.
        This is the single source of truth for table dimensions, using a hierarchical approach.
        """
        # --- Width Calculation ---
        width_emu = None
        try:
            # Priority 1: Standard way to get table width
            width_emu = table.width
            if width_emu:
                self.logger.debug("      - Found table width via 'table.width': %d EMU", width_emu)
        except AttributeError:
            width_emu = None

        if not width_emu:
            # Priority 2: Check cell width (handles cases where table width is not set)
            try:
                if table.rows and table.columns:
                    cell_width = table.rows[0].cells[0].width
                    if cell_width:
                        width_emu = cell_width
                        self.logger.debug("      - Found table width via 'cell.width': %d EMU", width_emu)
            except (AttributeError, IndexError):
                pass  # Cell width not available

        if not width_emu:
            # Priority 3: Low-level OXML check
            self.logger.debug("      - 'table.width' and 'cell.width' not found. Trying direct OXML access.")
            try:
                tblPr = table._element.tblPr
                if tblPr is not None:
                    tblW = tblPr.tblW
                    if tblW is not None:
                        width_emu = tblW.w
                        if width_emu:
                            self.logger.debug("      - Found table width via OXML 'tblW': %s", width_emu)
            except AttributeError:
                width_emu = None

        if not width_emu:
            # Priority 4: Fallback to page width
            section = doc.sections[-1]
            page_width_emu = section.page_width
            left_margin_emu = section.left_margin
            right_margin_emu = section.right_margin
            width_emu = page_width_emu - (left_margin_emu or 0) - (right_margin_emu or 0)
            self.logger.debug(
                "      - Table width not found, falling back to page width minus margins (%.2f\")",
                points_to_inches(emu_to_points(width_emu))
            )

        width_pts = emu_to_points(width_emu or 0)

        # --- Height Calculation ---
        # Sum of explicit row heights
        height_emu = sum(row.height for row in table.rows if row.height is not None)
        
        # Handle cases where all rows have auto height (a common scenario)
        if height_emu == 0 and table.rows:
            # Estimate height based on a standard row height
            # This is a fallback and might not be perfectly accurate
            estimated_row_height_pts = 14.4  # 0.2 inches, a reasonable default
            height_pts = len(table.rows) * estimated_row_height_pts
            self.logger.debug(
                "      - All table rows have auto height. Estimating height as %.2f\" for %d rows.",
                points_to_inches(height_pts), len(table.rows)
            )
        else:
            height_pts = emu_to_points(height_emu)
            # Log a warning if some rows are auto, as the height will be an underestimate
            if any(r.height is None for r in table.rows):
                self.logger.warning(
                    "      - Some table rows have auto height. Calculated height (%.2f\") may be an underestimate.",
                    points_to_inches(height_pts)
                )

        return {
            'width_pts': width_pts,
            'height_pts': height_pts,
            'width_inches': points_to_inches(width_pts),
            'height_inches': points_to_inches(height_pts),
        }

    def _process_table_placeholders(self, doc: Document, table_placeholders: List[Dict]) -> Dict[int, Dict[str, float]]:
        """Replace table placeholders with overlay markers or insert images directly."""
        self.logger.debug("  > Processing %d table placeholders...", len(table_placeholders))
        metadata = {}
        # Sort by index to process tables in document order
        sorted_placeholders = sorted(table_placeholders, key=lambda x: x['table_index'])

        for placeholder in sorted_placeholders:
            table_idx = placeholder['table_index']
            subtype = placeholder.get('subtype', 'overlay')  # Default to overlay for backward compatibility
            
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                
                if len(table.rows) != 1 or len(table.columns) != 1:
                    self.logger.warning(
                        "    - Skipping table %d as it is not a 1x1 table, which is required for overlays and images.", table_idx
                    )
                    continue

                if subtype == 'image':
                    self.logger.debug("    - Processing table %d for image insertion.", table_idx)
                    self._process_image_placeholder(table, placeholder, doc)
                else:  # overlay
                    self.logger.debug("    - Processing table %d for overlay.", table_idx)
                    self._process_overlay_placeholder(table, placeholder, table_idx, metadata, doc)
            else:
                self.logger.warning("    - Table index %d is out of bounds.", table_idx)

        return metadata

    def _process_image_placeholder(self, table, placeholder: Dict, doc: Document):
        """Process an IMAGE placeholder by inserting the image directly into the table cell."""
        # Use the resolved path that was computed during validation
        # This ensures relative paths are resolved relative to the document location
        file_path = placeholder.get('resolved_path') or placeholder.get('file_path')
        
        try:
            # Validate image file exists (should already be validated, but double-check)
            if not os.path.exists(file_path):
                self.logger.error("      - Image file not found: %s", file_path)
                return
            
            # Get table dimensions for image sizing
            dimensions = self._get_table_dimensions_in_points(table, doc)
            table_width_inches = dimensions['width_inches']
            table_height_inches = dimensions['height_inches']
            
            # Calculate image dimensions
            img_width, img_height = self._calculate_image_dimensions(
                file_path, placeholder, table_width_inches, table_height_inches
            )
            
            # Clear the cell and insert the image
            cell = table.cell(0, 0)
            cell.paragraphs[0].clear()
            
            # Add the image to the paragraph
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(file_path, width=Inches(img_width), height=Inches(img_height))
            
            self.logger.info("      - Successfully inserted image: %s (%.2f\" x %.2f\")", 
                           os.path.basename(file_path), img_width, img_height)
            
        except Exception as e:
            self.logger.error("      - Failed to insert image %s: %s", file_path, e)

    def _calculate_image_dimensions(self, file_path: str, placeholder: Dict, 
                                  table_width: float, table_height: float) -> tuple:
        """Calculate appropriate image dimensions based on parameters and table size."""
        # Get explicit width/height from placeholder parameters
        width_param = placeholder.get('width')
        height_param = placeholder.get('height')
        
        try:
            # Get original image dimensions
            with Image.open(file_path) as img:
                original_width, original_height = img.size
                aspect_ratio = original_width / original_height
                
            # Convert to inches (assuming 96 DPI)
            original_width_inches = original_width / 96
            original_height_inches = original_height / 96
            
            # Apply user-specified dimensions or auto-fit to table
            if width_param and height_param:
                # Both dimensions specified
                try:
                    img_width = float(width_param.rstrip('in"'))
                    img_height = float(height_param.rstrip('in"'))
                except ValueError:
                    # If parsing fails, fall back to auto-sizing
                    img_width, img_height = self._auto_size_image(
                        original_width_inches, original_height_inches, 
                        table_width, table_height, aspect_ratio
                    )
            elif width_param:
                # Only width specified, maintain aspect ratio
                try:
                    img_width = float(width_param.rstrip('in"'))
                    img_height = img_width / aspect_ratio
                except ValueError:
                    img_width, img_height = self._auto_size_image(
                        original_width_inches, original_height_inches, 
                        table_width, table_height, aspect_ratio
                    )
            elif height_param:
                # Only height specified, maintain aspect ratio
                try:
                    img_height = float(height_param.rstrip('in"'))
                    img_width = img_height * aspect_ratio
                except ValueError:
                    img_width, img_height = self._auto_size_image(
                        original_width_inches, original_height_inches, 
                        table_width, table_height, aspect_ratio
                    )
            else:
                # No dimensions specified, auto-fit to table
                img_width, img_height = self._auto_size_image(
                    original_width_inches, original_height_inches, 
                    table_width, table_height, aspect_ratio
                )
                
            return img_width, img_height
            
        except Exception as e:
            self.logger.warning("      - Could not determine image dimensions for %s: %s", file_path, e)
            # Fallback to table dimensions with some padding
            return max(0.5, table_width * 0.9), max(0.5, table_height * 0.9)

    def _auto_size_image(self, original_width: float, original_height: float,
                        table_width: float, table_height: float, aspect_ratio: float) -> tuple:
        """Automatically size image to fit within table while maintaining aspect ratio."""
        # Add some padding (10% smaller than table)
        max_width = table_width * 0.9
        max_height = table_height * 0.9
        
        # Scale to fit within bounds while maintaining aspect ratio
        width_scale = max_width / original_width if original_width > 0 else 1
        height_scale = max_height / original_height if original_height > 0 else 1
        scale = min(width_scale, height_scale, 1.0)  # Don't upscale
        
        return original_width * scale, original_height * scale

    def _process_overlay_placeholder(self, table, placeholder: Dict, table_idx: int, 
                                   metadata: Dict, doc: Document):
        """Process an OVERLAY placeholder by inserting markers for PDF processing."""
        dimensions = self._get_table_dimensions_in_points(table, doc)
        metadata[table_idx] = dimensions
        self.logger.debug("      - Recorded table %d dimensions: %.2f\" x %.2f\" (WxH)",
                         table_idx, dimensions['width_inches'], dimensions['height_inches'])

        # Clear the cell and place the primary marker
        primary_cell = table.cell(0, 0)
        marker = Config.get_overlay_marker(table_idx, page_num=1)
        primary_cell.paragraphs[0].clear()
        primary_cell.paragraphs[0].add_run(marker)
        self.logger.debug("      - Placed primary marker in table %d: %s", table_idx, marker)

        # Replicate rows for multi-page overlays
        # Calculate actual number of pages that will be selected based on page specification
        total_pages_in_source = placeholder.get('page_count', 1)
        page_spec = placeholder.get('page_spec')
        
        self.logger.debug("      - Calculating pages for spec '%s', total source pages: %d", page_spec, total_pages_in_source)
        
        if page_spec:
            try:
                # Import PageSelector to calculate actual selected pages
                from ..utils.page_selector import PageSelector
                page_selector = PageSelector()
                
                # Create a mock document with the correct page count to test selection
                class MockDocument:
                    def __init__(self, page_count):
                        self._page_count = page_count
                    def __len__(self):
                        return self._page_count
                
                mock_doc = MockDocument(total_pages_in_source)
                page_selection = page_selector.parse_specification(page_spec)
                selected_pages = page_selector.apply_selection(mock_doc, page_selection)
                num_pages = len(selected_pages)
                
                self.logger.debug("      - Page spec '%s' selects %d of %d pages", page_spec, num_pages, total_pages_in_source)
            except Exception as e:
                self.logger.error("      - Error calculating page selection: %s", e)
                num_pages = total_pages_in_source
        else:
            num_pages = total_pages_in_source
        
        self.logger.debug("      - Final num_pages: %d", num_pages)
        
        if num_pages > 1:
            self._replicate_table_rows_for_overlay(table, num_pages, table_idx)
