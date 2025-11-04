#!/usr/bin/env python3
"""
Flask Web Interface for PyArchInit-Mini
"""

import os
from pathlib import Path
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from flask_login import login_required, current_user
from flask_socketio import SocketIO
from flask_babel import lazy_gettext as _l
from wtforms import StringField, TextAreaField, IntegerField, SelectField, FileField, BooleanField
from wtforms.validators import DataRequired, Optional
from werkzeug.utils import secure_filename
import tempfile
import base64
from sqlalchemy import text

# PyArchInit-Mini imports
from pyarchinit_mini import __version__
from pyarchinit_mini.database.connection import DatabaseConnection
from pyarchinit_mini.database.manager import DatabaseManager
from pyarchinit_mini.services.site_service import SiteService
from pyarchinit_mini.services.us_service import USService
from pyarchinit_mini.services.inventario_service import InventarioService
from pyarchinit_mini.services.thesaurus_service import ThesaurusService
from pyarchinit_mini.services.user_service import UserService
from pyarchinit_mini.services.analytics_service import AnalyticsService
from pyarchinit_mini.services.relationship_sync_service import RelationshipSyncService
from pyarchinit_mini.services.datazione_service import DatazioneService
from pyarchinit_mini.services.media_service import MediaService
from pyarchinit_mini.services.import_export_service import ImportExportService
from pyarchinit_mini.harris_matrix.matrix_generator import HarrisMatrixGenerator
from pyarchinit_mini.harris_matrix.matrix_visualizer import MatrixVisualizer
from pyarchinit_mini.harris_matrix.pyarchinit_visualizer import PyArchInitMatrixVisualizer
from pyarchinit_mini.utils.stratigraphic_validator import StratigraphicValidator
from pyarchinit_mini.pdf_export.pdf_generator import PDFGenerator
from pyarchinit_mini.media_manager.media_handler import MediaHandler
from pyarchinit_mini.graphml_converter import convert_dot_content_to_graphml

# Import authentication routes
from pyarchinit_mini.web_interface.auth_routes import auth_bp, init_login_manager, write_permission_required

# Import PyArchInit import/export routes
from pyarchinit_mini.web_interface.pyarchinit_import_export_routes import pyarchinit_import_export_bp

# Import Harris Matrix Creator routes
from pyarchinit_mini.web_interface.harris_creator_routes import harris_creator_bp

# Import Excel Import routes
from pyarchinit_mini.web_interface.excel_import_routes import excel_import_bp

# Import EM Node Configuration routes
from pyarchinit_mini.web_interface.em_node_config_routes import em_node_config_bp

# Import 3D Builder routes
from pyarchinit_mini.web_interface.three_d_builder_routes import three_d_builder_bp, three_d_builder_ui_bp

# Import WebSocket events
from pyarchinit_mini.web_interface.socketio_events import (
    init_socketio_events,
    init_blender_socketio_events,
    broadcast_site_created,
    broadcast_site_updated,
    broadcast_site_deleted,
    broadcast_us_created,
    broadcast_us_updated,
    broadcast_us_deleted,
    broadcast_inventario_created,
    broadcast_inventario_updated,
    broadcast_inventario_deleted,
    broadcast_blender_command
)

# Forms
class SiteForm(FlaskForm):
    sito = StringField(_l('Site Name'), validators=[DataRequired()])
    nazione = StringField(_l('Country'))
    regione = StringField(_l('Region'))
    comune = StringField(_l('Municipality'))
    provincia = StringField(_l('Province'))
    definizione_sito = StringField(_l('Site Definition'))
    descrizione = TextAreaField(_l('Description'))

class USForm(FlaskForm):
    # TAB 1: Informazioni Base
    # Identificazione
    sito = SelectField(_l('Site'), validators=[DataRequired()], coerce=str)
    area = StringField(_l('Area'))  # Already a StringField
    us = StringField(_l('US Number'), validators=[DataRequired()])  # Changed from IntegerField to StringField
    unita_tipo = SelectField(_l('Unit Type'), choices=[
        ('', _l('-- Select --')),
        ('US', _l('US')),
        ('USM', _l('USM')),
        ('VSF', _l('VSF')),
        ('SF', _l('SF')),
        ('CON', _l('CON')),
        ('USD', _l('USD')),
        ('USVA', _l('USVA')),
        ('USVB', _l('USVB')),
        ('USVC', _l('USVC')),
        ('DOC', _l('DOC')),
        ('TU', _l('TU')),
        ('property', _l('property')),
        ('Combiner', _l('Combiner')),
        ('Extractor', _l('Extractor'))
    ])
    tipo_documento = SelectField(_l('Document Type'), choices=[
        ('', _l('-- Select --')),
        ('image', _l('Image')),
        ('PDF', _l('PDF')),
        ('DOCX', _l('DOCX')),
        ('CSV', _l('CSV')),
        ('Excel', _l('Excel')),
        ('TXT', _l('TXT'))
    ])
    documento_file = FileField(_l('Upload Document File'))

    # Dati di Scavo
    anno_scavo = IntegerField(_l('Excavation Year'), validators=[Optional()])
    scavato = SelectField(_l('Excavated'), choices=[
        ('', _l('-- Select --')),
        ('Sì', _l('Yes')),
        ('No', _l('No')),
        ('Parzialmente', _l('Partially'))
    ])
    schedatore = StringField(_l('Cataloguer'))
    metodo_di_scavo = SelectField(_l('Excavation Method'), choices=[
        ('', _l('-- Select --')),
        ('Manuale', _l('Manual')),
        ('Meccanico', _l('Mechanical')),
        ('Misto', _l('Mixed'))
    ])
    data_schedatura = StringField(_l('Cataloguing Date'), description=_l('Format: YYYY-MM-DD'))
    attivita = StringField(_l('Activity'))

    # Responsabili
    direttore_us = StringField(_l('US Director'))
    responsabile_us = StringField(_l('US Supervisor'))

    # Contesto
    settore = StringField(_l('Sector'))
    quad_par = StringField(_l('Square/Partition'))
    ambient = StringField(_l('Room'))
    saggio = StringField(_l('Trench'))

    # Catalogazione ICCD
    n_catalogo_generale = StringField(_l('General Catalogue No.'))
    n_catalogo_interno = StringField(_l('Internal Catalogue No.'))
    n_catalogo_internazionale = StringField(_l('International Catalogue No.'))
    soprintendenza = StringField(_l('Superintendency'))

    # TAB 2: Descrizioni
    d_stratigrafica = TextAreaField(_l('Stratigraphic Description'))
    d_interpretativa = TextAreaField(_l('Interpretative Description'))
    descrizione = TextAreaField(_l('Detailed Description'))
    interpretazione = TextAreaField(_l('Interpretation'))
    osservazioni = TextAreaField(_l('Observations'))

    # TAB 3: Caratteristiche Fisiche
    definizione_stratigrafica = SelectField('Definizione Stratigrafica', choices=[], coerce=str)
    formazione = SelectField(_l('Formation'), choices=[], coerce=str)
    stato_di_conservazione = SelectField(_l('Conservation State'), choices=[
        ('', _l('-- Select --')),
        ('Ottimo', _l('Excellent')),
        ('Buono', _l('Good')),
        ('Discreto', _l('Fair')),
        ('Cattivo', _l('Poor'))
    ])
    colore = SelectField(_l('Color'), choices=[], coerce=str)
    consistenza = SelectField('Consistenza', choices=[], coerce=str)
    struttura = StringField('Struttura')

    # Misure
    quota_relativa = StringField('Quota Relativa')
    quota_abs = StringField('Quota Assoluta')
    lunghezza_max = StringField('Lunghezza Max (cm)')
    larghezza_media = StringField('Larghezza Media (cm)')
    altezza_max = StringField('Altezza Max (cm)')
    altezza_min = StringField('Altezza Min (cm)')
    profondita_max = StringField('Profondità Max (cm)')
    profondita_min = StringField('Profondità Min (cm)')

    # TAB 4: Cronologia
    periodo_iniziale = StringField('Periodo Iniziale')
    fase_iniziale = StringField('Fase Iniziale')
    periodo_finale = StringField('Periodo Finale')
    fase_finale = StringField('Fase Finale')
    datazione = SelectField('Datazione', choices=[], coerce=str)
    affidabilita = SelectField('Affidabilità', choices=[
        ('', '-- Seleziona --'),
        ('Alta', 'Alta'),
        ('Media', 'Media'),
        ('Bassa', 'Bassa')
    ])

    # TAB 5: Relazioni Stratigrafiche
    rapporti = TextAreaField('Rapporti Stratigrafici',
                            description='Formato: copre 1002, taglia 1005, si appoggia a 1010')

    # TAB 6: Documentazione
    inclusi = TextAreaField('Inclusi')
    campioni = TextAreaField('Campioni')
    documentazione = TextAreaField('Documentazione')
    cont_per = TextAreaField('Contenitori/Contenuti')

    # Altri campi
    flottazione = SelectField('Flottazione', choices=[
        ('', '-- Seleziona --'),
        ('Sì', 'Sì'),
        ('No', 'No')
    ])
    setacciatura = SelectField('Setacciatura', choices=[
        ('', '-- Seleziona --'),
        ('Sì', 'Sì'),
        ('No', 'No')
    ])

class InventarioForm(FlaskForm):
    # TAB 1: Identificazione
    sito = SelectField('Sito', validators=[DataRequired()], coerce=str)
    numero_inventario = IntegerField('Numero Inventario', validators=[DataRequired()])
    n_reperto = IntegerField('N. Reperto', validators=[Optional()])
    schedatore = StringField('Schedatore')
    date_scheda = StringField('Data Scheda', description='Formato: AAAA-MM-GG')
    years = IntegerField('Anno', validators=[Optional()])

    # TAB 2: Classificazione
    tipo_reperto = SelectField('Tipo Reperto', choices=[])  # Populated from thesaurus
    criterio_schedatura = StringField('Criterio Schedatura')
    definizione = StringField('Definizione')
    tipo = StringField('Tipo')
    tipo_contenitore = StringField('Tipo Contenitore')
    struttura = StringField('Struttura')
    descrizione = TextAreaField('Descrizione')

    # TAB 3: Contesto
    area = StringField('Area')
    us = StringField('US')  # Text field as per model
    punto_rinv = StringField('Punto Rinvenimento')
    elementi_reperto = TextAreaField('Elementi Reperto')

    # TAB 4: Caratteristiche Fisiche
    stato_conservazione = SelectField('Stato di Conservazione', choices=[])  # From thesaurus
    lavato = SelectField('Lavato', choices=[
        ('', '-- Seleziona --'),
        ('Sì', 'Sì'),
        ('No', 'No')
    ])
    nr_cassa = StringField('N. Cassa')
    luogo_conservazione = StringField('Luogo di Conservazione')

    # TAB 5: Conservazione e Gestione
    repertato = SelectField('Repertato', choices=[
        ('', '-- Seleziona --'),
        ('Sì', 'Sì'),
        ('No', 'No')
    ])
    diagnostico = SelectField('Diagnostico', choices=[
        ('', '-- Seleziona --'),
        ('Sì', 'Sì'),
        ('No', 'No')
    ])

    # TAB 6: Caratteristiche Ceramiche
    corpo_ceramico = SelectField('Corpo Ceramico', choices=[])  # From thesaurus
    rivestimento = SelectField('Rivestimento', choices=[])  # From thesaurus
    diametro_orlo = StringField('Diametro Orlo (cm)')
    eve_orlo = StringField('EVE Orlo')

    # TAB 7: Misurazioni
    peso = StringField('Peso (g)')
    forme_minime = IntegerField('Forme Minime', validators=[Optional()])
    forme_massime = IntegerField('Forme Massime', validators=[Optional()])
    totale_frammenti = IntegerField('Totale Frammenti', validators=[Optional()])
    misurazioni = TextAreaField('Misurazioni')

    # TAB 8: Documentazione
    datazione_reperto = StringField('Datazione Reperto')
    rif_biblio = TextAreaField('Riferimenti Bibliografici')
    tecnologie = TextAreaField('Tecnologie')
    negativo_photo = StringField('Negativo Fotografico')
    diapositiva = StringField('Diapositiva')

class MediaUploadForm(FlaskForm):
    entity_type = SelectField('Entity Type', choices=[
        ('', '-- Select --'),
        ('site', 'Site'),
        ('us', 'US'),
        ('inventario', 'Inventory')
    ], validators=[DataRequired()])

    # Fields for Site selection
    site_id = SelectField('Site', choices=[], coerce=int, validators=[Optional()])

    # Fields for US selection
    us_site = SelectField('Site', choices=[], coerce=str, validators=[Optional()])
    us_area = StringField('Area', validators=[Optional()])
    us_number = StringField('US Number', validators=[Optional()])

    # Fields for Inventory selection
    inv_site = SelectField('Site', choices=[], coerce=str, validators=[Optional()])
    inv_number = StringField('Inventory Number', validators=[Optional()])

    file = FileField('File', validators=[DataRequired()])
    description = TextAreaField('Description')
    author = StringField('Author/Photographer')

class DatabaseUploadForm(FlaskForm):
    """Form for uploading SQLite database files"""
    database_file = FileField('Database SQLite (.db)', validators=[DataRequired()])
    database_name = StringField('Nome Database', validators=[DataRequired()],
                               description='Nome identificativo per questo database')
    description = TextAreaField('Descrizione',
                               description='Descrizione opzionale del database')

class DatabaseConnectionForm(FlaskForm):
    """Form for PostgreSQL database connections"""
    db_type = SelectField('Tipo Database', choices=[
        ('postgresql', 'PostgreSQL'),
        ('sqlite', 'SQLite (file locale)')
    ], validators=[DataRequired()])

    # PostgreSQL fields
    host = StringField('Host', description='Es: localhost, 192.168.1.100')
    port = IntegerField('Porta', description='Default: 5432 (PostgreSQL)')
    database = StringField('Nome Database', validators=[DataRequired()])
    username = StringField('Username')
    password = StringField('Password')

    # SQLite fields
    sqlite_path = StringField('Percorso File SQLite',
                             description='Es: /path/to/database.db')

    connection_name = StringField('Nome Connessione', validators=[DataRequired()],
                                 description='Nome identificativo per questa connessione')

class GraphMLExportForm(FlaskForm):
    """Form for GraphML export of Harris Matrix"""
    site = SelectField('Sito', validators=[DataRequired()], coerce=str)
    title = StringField('Titolo Diagramma', description='Intestazione opzionale per il diagramma')
    grouping = SelectField('Raggruppamento', choices=[
        ('period_area', 'Periodo + Area'),
        ('period', 'Solo Periodo'),
        ('area', 'Solo Area'),
        ('none', 'Nessun Raggruppamento')
    ], default='period_area')
    reverse_epochs = BooleanField('Inverti ordine periodi', default=False)

# Flask App Setup
def create_app():
    # Declare global variables for database and services
    # This allows switch_database() to access and modify them
    global db_conn, db_manager
    global site_service, us_service, inventario_service, thesaurus_service
    global user_service, analytics_service, relationship_sync_service, datazione_service
    global matrix_generator, export_import_service, media_service

    app = Flask(__name__)
    app.config['SECRET_KEY'] = 'your-secret-key-here'
    # Use centralized ~/.pyarchinit_mini directory
    pyarchinit_home = Path.home() / '.pyarchinit_mini'
    app.config['UPLOAD_FOLDER'] = str(pyarchinit_home / 'web_interface' / 'static' / 'uploads')
    app.config['DATABASE_FOLDER'] = str(pyarchinit_home / 'data')

    # Initialize CSRF protection
    csrf = CSRFProtect(app)

    # Initialize Flask-Babel for i18n
    from pyarchinit_mini.i18n import init_babel, get_locale
    from flask_babel import gettext as _
    babel = init_babel(app)

    # Make get_locale and gettext available in all templates
    @app.context_processor
    def inject_locale():
        return dict(get_locale=get_locale, _=_, version=__version__)

    # Create necessary folders
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['DATABASE_FOLDER'], exist_ok=True)

    # Initialize database
    # Use centralized ~/.pyarchinit_mini/data directory
    default_db_path = str(pyarchinit_home / 'data' / 'pyarchinit_mini.db')
    default_db_url = f"sqlite:///{default_db_path}"
    database_url = os.getenv("DATABASE_URL", default_db_url)
    print(f"[FLASK] Current working directory: {os.getcwd()}")
    print(f"[FLASK] Using database: {database_url}")
    # If SQLite, show absolute path
    if database_url.startswith('sqlite:///'):
        db_path = database_url.replace('sqlite:///', '')
        if not db_path.startswith('/'):
            abs_path = os.path.abspath(db_path)
            print(f"[FLASK] SQLite absolute path: {abs_path}")
            print(f"[FLASK] Database exists: {os.path.exists(abs_path)}")
    db_conn = DatabaseConnection.from_url(database_url)
    db_conn.create_tables()
    db_manager = DatabaseManager(db_conn)
    print(f"[FLASK] Database connection initialized")

    # Store current database info in app config
    app.config['CURRENT_DATABASE_URL'] = database_url
    app.config['DATABASE_CONNECTIONS'] = {}  # Store named connections

    # Add default database connections for easy switching
    project_root = os.path.dirname(os.path.dirname(__file__))
    root_db_path = os.path.join(project_root, 'pyarchinit_mini.db')
    package_db_path = os.path.join(os.path.dirname(__file__), '..', 'pyarchinit_mini', 'pyarchinit_mini.db')
    package_db_path = os.path.abspath(package_db_path)

    app.config['DATABASE_CONNECTIONS']['Database Root (Progetto)'] = {
        'type': 'sqlite',
        'path': root_db_path,
        'url': f'sqlite:///{root_db_path}',
        'description': 'Database nella cartella root del progetto (per sviluppo PyCharm)',
        'uploaded': False,
        'is_active': database_url == f'sqlite:///{root_db_path}'
    }

    app.config['DATABASE_CONNECTIONS']['Database Package'] = {
        'type': 'sqlite',
        'path': package_db_path,
        'url': f'sqlite:///{package_db_path}',
        'description': 'Database nella cartella del package pyarchinit_mini',
        'uploaded': False,
        'is_active': database_url == f'sqlite:///{package_db_path}'
    }

    # Save active default database to ConnectionManager for persistence
    from pyarchinit_mini.config.connection_manager import get_connection_manager
    conn_manager = get_connection_manager()

    # Find and save the active default database
    for db_name, db_config in app.config['DATABASE_CONNECTIONS'].items():
        if db_config.get('is_active', False) and not db_config.get('uploaded', False):
            # Check if already exists before adding
            existing = conn_manager.get_connection(db_name)
            if not existing:
                conn_manager.add_connection(
                    name=db_name,
                    db_type=db_config['type'],
                    connection_string=db_config['url'],
                    description=db_config.get('description', 'Default database')
                )
                print(f"[FLASK] Added default database '{db_name}' to saved connections")
            break

    # Initialize services
    site_service = SiteService(db_manager)
    us_service = USService(db_manager)
    inventario_service = InventarioService(db_manager)
    thesaurus_service = ThesaurusService(db_manager)
    user_service = UserService(db_manager)
    analytics_service = AnalyticsService(db_manager)
    relationship_sync_service = RelationshipSyncService(db_manager)
    datazione_service = DatazioneService(db_manager)
    matrix_generator = HarrisMatrixGenerator(db_manager, us_service)  # Pass us_service for proper matrix generation
    export_import_service = ImportExportService(db_manager.connection.connection_string)
    matrix_visualizer = MatrixVisualizer()
    graphviz_visualizer = PyArchInitMatrixVisualizer()  # Graphviz visualizer (desktop GUI style)
    pdf_generator = PDFGenerator()
    media_handler = MediaHandler()
    media_service = MediaService(db_manager, media_handler)

    # Store services in app for access in routes
    app.db_manager = db_manager
    app.user_service = user_service
    app.site_service = site_service
    app.us_service = us_service
    app.inventario_service = inventario_service
    app.thesaurus_service = thesaurus_service
    app.analytics_service = analytics_service
    app.relationship_sync_service = relationship_sync_service
    app.datazione_service = datazione_service
    app.matrix_generator = matrix_generator
    app.export_import_service = export_import_service
    app.media_service = media_service

    # Initialize Flask-Login
    init_login_manager(app, user_service)

    # Register authentication blueprint
    app.register_blueprint(auth_bp)

    # Register PyArchInit import/export blueprint
    app.register_blueprint(pyarchinit_import_export_bp, url_prefix='/pyarchinit-import-export')

    # Register Harris Matrix Creator blueprint
    app.register_blueprint(harris_creator_bp)

    # Register Excel Import blueprint
    app.register_blueprint(excel_import_bp, url_prefix='/excel-import')

    # Register EM Node Configuration blueprint
    app.register_blueprint(em_node_config_bp)

    # Register 3D Builder blueprints
    app.register_blueprint(three_d_builder_bp)  # API routes
    app.register_blueprint(three_d_builder_ui_bp)  # UI routes

    # Exempt PyArchInit API endpoints from CSRF protection (JSON APIs)
    csrf.exempt(pyarchinit_import_export_bp)
    csrf.exempt(harris_creator_bp)
    csrf.exempt(three_d_builder_bp)
    csrf.exempt(excel_import_bp)
    csrf.exempt(em_node_config_bp)

    # Initialize Flask-SocketIO
    socketio = SocketIO(app, cors_allowed_origins="*")

    # Initialize WebSocket event handlers
    init_socketio_events(socketio)
    init_blender_socketio_events(socketio)  # Real-time Blender streaming

    # Store socketio in app for access in routes
    app.socketio = socketio

    # Helper function to get thesaurus values
    def get_thesaurus_choices(field_name, table_name='inventario_materiali_table'):
        """Get thesaurus choices for a field"""
        try:
            values = thesaurus_service.get_field_values(table_name, field_name)
            return [('', '-- Seleziona --')] + [(v['value'], v['label']) for v in values]
        except Exception:
            # Return empty list if thesaurus not available
            return [('', '-- Seleziona --')]
    
    # Routes
    @app.route('/')
    @login_required
    def index():
        """Dashboard with statistics"""
        try:
            # Get basic statistics
            sites = site_service.get_all_sites(size=5)
            total_sites = site_service.count_sites()
            total_us = us_service.count_us()
            total_inventory = inventario_service.count_inventario()
            
            stats = {
                'total_sites': total_sites,
                'total_us': total_us,
                'total_inventory': total_inventory,
                'recent_sites': sites
            }
            
            return render_template('dashboard.html', stats=stats)
        except Exception as e:
            flash(f'Errore caricamento dashboard: {str(e)}', 'error')
            return render_template('dashboard.html', stats={})

    @app.route('/analytics')
    @login_required
    def analytics():
        """Analytics dashboard with charts"""
        try:
            # Get all analytics data
            analytics_data = analytics_service.get_complete_dashboard_data()

            return render_template('analytics/dashboard.html', data=analytics_data)
        except Exception as e:
            flash(f'Errore caricamento analytics: {str(e)}', 'error')
            return redirect(url_for('index'))

    # Sites routes
    @app.route('/sites')
    @login_required
    def sites_list():
        page = request.args.get('page', 1, type=int)
        search = request.args.get('search', '')
        
        if search:
            sites = site_service.search_sites(search, page=page, size=20)
        else:
            sites = site_service.get_all_sites(page=page, size=20)
        
        total = site_service.count_sites()
        
        return render_template('sites/list.html', sites=sites, total=total, 
                             page=page, search=search)
    
    @app.route('/sites/create', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def create_site():
        form = SiteForm()
        
        if form.validate_on_submit():
            try:
                site_data = {
                    'sito': form.sito.data,
                    'nazione': form.nazione.data,
                    'regione': form.regione.data,
                    'comune': form.comune.data,
                    'provincia': form.provincia.data,
                    'definizione_sito': form.definizione_sito.data,
                    'descrizione': form.descrizione.data
                }
                
                site = site_service.create_site(site_data)

                # Try to get site ID, use None if instance is detached
                try:
                    site_id = site.id_sito
                except Exception:
                    site_id = None

                # Broadcast site creation
                broadcast_site_created(socketio, site_data['sito'], site_id)

                flash(f'Sito "{site_data["sito"]}" creato con successo!', 'success')
                return redirect(url_for('sites_list'))
                
            except Exception as e:
                flash(f'Errore nella creazione del sito: {str(e)}', 'error')
        
        return render_template('sites/form.html', form=form, title='New Site')

    @app.route('/sites/<int:site_id>/edit', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def edit_site(site_id):
        """Edit existing site"""
        form = SiteForm()

        # Get existing site
        site = site_service.get_site_dto_by_id(site_id)
        if not site:
            flash('Sito non trovato', 'error')
            return redirect(url_for('sites_list'))

        if form.validate_on_submit():
            try:
                update_data = {
                    'sito': form.sito.data,
                    'nazione': form.nazione.data,
                    'regione': form.regione.data,
                    'comune': form.comune.data,
                    'provincia': form.provincia.data,
                    'definizione_sito': form.definizione_sito.data,
                    'descrizione': form.descrizione.data
                }

                updated_site = site_service.update_site(site_id, update_data)

                flash(f'Sito "{update_data["sito"]}" aggiornato con successo!', 'success')
                return redirect(url_for('sites_list'))

            except Exception as e:
                flash(f'Errore nell\'aggiornamento del sito: {str(e)}', 'error')

        # Pre-populate form with existing data
        elif request.method == 'GET':
            form.sito.data = site.sito
            form.nazione.data = site.nazione
            form.regione.data = site.regione
            form.comune.data = site.comune
            form.provincia.data = site.provincia
            form.definizione_sito.data = site.definizione_sito
            form.descrizione.data = site.descrizione

        # Get associated media files
        try:
            media_list = media_service.get_media_by_entity('site', site_id, size=100)
        except Exception as e:
            logger.warning(f"Failed to fetch media for site {site_id}: {e}")
            media_list = []

        return render_template('sites/form.html', form=form, title='Edit Site', edit_mode=True,
                             site_id=site_id, media_list=media_list)

    @app.route('/sites/<int:site_id>/delete', methods=['POST'])
    @login_required
    @write_permission_required
    def delete_site(site_id):
        """Delete site and all related records (cascade deletion)"""
        try:
            # Get site name before deletion for flash message
            site = site_service.get_site_dto_by_id(site_id)
            if not site:
                return jsonify({'success': False, 'message': 'Sito non trovato'}), 404

            site_name = site.sito

            # Perform cascade deletion
            deletion_stats = site_service.delete_site(site_id)

            # Build success message
            message = f'Sito "{site_name}" eliminato con successo!'
            if deletion_stats['us'] > 0:
                message += f" ({deletion_stats['us']} US, {deletion_stats['inventario']} inventario, "
                message += f"{deletion_stats['us_relationships']} relazioni, "
                message += f"{deletion_stats['periodizzazione']} periodizzazioni eliminate)"

            return jsonify({
                'success': True,
                'message': message,
                'stats': deletion_stats
            })

        except Exception as e:
            # Flask will automatically log the exception traceback
            return jsonify({
                'success': False,
                'message': f'Errore durante l\'eliminazione del sito: {str(e)}'
            }), 500

    @app.route('/sites/<int:site_id>')
    @login_required
    def view_site(site_id):
        # Get site and related data within session scope
        with db_manager.connection.get_session() as session:
            from pyarchinit_mini.models.site import Site as SiteModel
            from pyarchinit_mini.models.us import US as USModel
            from pyarchinit_mini.models.inventario_materiali import InventarioMateriali as InvModel

            site = session.query(SiteModel).filter(SiteModel.id_sito == site_id).first()
            if not site:
                flash('Sito non trovato', 'error')
                return redirect(url_for('sites_list'))

            site_name = site.sito

            # Get related data and convert to dicts within session
            us_records = session.query(USModel).filter(USModel.sito == site_name).limit(50).all()
            us_list = [us.to_dict() for us in us_records]

            inv_records = session.query(InvModel).filter(InvModel.sito == site_name).limit(50).all()
            inventory_list = [inv.to_dict() for inv in inv_records]

            # Convert site to dict
            site_dict = site.to_dict()

        # Use dicts outside session
        return render_template('sites/detail.html', site=site_dict,
                             us_list=us_list, inventory_list=inventory_list)
    
    # US routes
    @app.route('/us')
    @login_required
    def us_list():
        from flask import session

        page = request.args.get('page', 1, type=int)

        # Advanced filters
        sito_filter = request.args.get('sito', '')
        area_filter = request.args.get('area', '')
        unita_tipo_filter = request.args.get('unita_tipo', '')
        anno_scavo_filter = request.args.get('anno_scavo', '')
        periodo_filter = request.args.get('periodo', '')
        fase_filter = request.args.get('fase', '')
        us_number_filter = request.args.get('us_number', '')

        # Build filters dictionary
        filters = {}
        if sito_filter:
            filters['sito'] = sito_filter
        if area_filter:
            filters['area'] = area_filter
        if unita_tipo_filter:
            filters['unita_tipo'] = unita_tipo_filter
        if anno_scavo_filter:
            filters['anno_scavo'] = int(anno_scavo_filter)
        if us_number_filter:
            filters['us'] = us_number_filter

        # Save filters in session for navigation and PDF export
        session['us_filters'] = {
            'sito': sito_filter,
            'area': area_filter,
            'unita_tipo': unita_tipo_filter,
            'anno_scavo': anno_scavo_filter,
            'periodo': periodo_filter,
            'fase': fase_filter,
            'us_number': us_number_filter
        }

        us_list = us_service.get_all_us(page=page, size=20, filters=filters)
        total = us_service.count_us(filters=filters)

        # Get sites for filter
        sites = site_service.get_all_sites(size=100)

        # Get distinct values for filters
        with db_manager.connection.get_session() as db_session:
            from pyarchinit_mini.models.us import US as USModel
            from sqlalchemy import distinct

            areas = db_session.query(distinct(USModel.area)).filter(USModel.area.isnot(None), USModel.area != '').all()
            areas = sorted([a[0] for a in areas if a[0]])

            unit_types = ['US', 'USM', 'VSF', 'SF', 'CON', 'USD', 'USVA', 'USVB', 'USVC', 'DOC', 'TU', 'property', 'Combiner', 'Extractor']

            years = db_session.query(distinct(USModel.anno_scavo)).filter(USModel.anno_scavo.isnot(None)).all()
            years = sorted([y[0] for y in years if y[0]], reverse=True)

        return render_template('us/list.html',
                             us_list=us_list,
                             sites=sites,
                             areas=areas,
                             unit_types=unit_types,
                             years=years,
                             total=total,
                             page=page,
                             sito_filter=sito_filter,
                             area_filter=area_filter,
                             unita_tipo_filter=unita_tipo_filter,
                             anno_scavo_filter=anno_scavo_filter,
                             periodo_filter=periodo_filter,
                             fase_filter=fase_filter,
                             us_number_filter=us_number_filter)
    
    @app.route('/us/create', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def create_us():
        form = USForm()
        
        # Populate site choices
        sites = site_service.get_all_sites(size=100)
        form.sito.choices = [('', '-- Seleziona Sito --')] + [(s.sito, s.sito) for s in sites]

        # Populate datazione choices
        datazioni_choices = datazione_service.get_datazioni_choices()
        form.datazione.choices = [('', '-- Seleziona Datazione --')] + [(d['value'], d['label']) for d in datazioni_choices]

        # Populate thesaurus choices for US fields
        try:
            def_strat_values = thesaurus_service.get_field_values('us_table', 'definizione_stratigrafica')
            form.definizione_stratigrafica.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in def_strat_values]

            formazione_values = thesaurus_service.get_field_values('us_table', 'formazione')
            form.formazione.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in formazione_values]

            consistenza_values = thesaurus_service.get_field_values('us_table', 'consistenza')
            form.consistenza.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in consistenza_values]

            colore_values = thesaurus_service.get_field_values('us_table', 'colore')
            form.colore.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in colore_values]
        except Exception as e:
            print(f"Error loading thesaurus values: {e}")
            # Fallback to empty choices
            form.definizione_stratigrafica.choices = [('', '-- Seleziona --')]
            form.formazione.choices = [('', '-- Seleziona --')]
            form.consistenza.choices = [('', '-- Seleziona --')]
            form.colore.choices = [('', '-- Seleziona --')]

        if form.validate_on_submit():
            try:
                # Helper function to convert numeric fields
                def to_float(value):
                    if value and str(value).strip():
                        try:
                            return float(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                def to_int(value):
                    if value and str(value).strip():
                        try:
                            return int(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                def to_date(value):
                    if value and str(value).strip():
                        try:
                            from datetime import datetime
                            # Try parsing date string (YYYY-MM-DD)
                            return datetime.strptime(str(value).strip(), '%Y-%m-%d').date()
                        except (ValueError, TypeError):
                            return None
                    return None

                us_data = {
                    # TAB 1: Informazioni Base
                    'sito': form.sito.data,
                    'area': form.area.data,
                    'us': to_int(form.us.data),
                    'unita_tipo': form.unita_tipo.data,
                    'tipo_documento': form.tipo_documento.data if form.unita_tipo.data == 'DOC' else None,
                    'anno_scavo': form.anno_scavo.data,
                    'scavato': form.scavato.data,
                    'schedatore': form.schedatore.data,
                    'metodo_di_scavo': form.metodo_di_scavo.data,
                    'data_schedatura': to_date(form.data_schedatura.data),
                    'attivita': form.attivita.data,
                    'direttore_us': form.direttore_us.data,
                    'responsabile_us': form.responsabile_us.data,
                    'settore': form.settore.data,
                    'quad_par': form.quad_par.data,
                    'ambient': form.ambient.data,
                    'saggio': form.saggio.data,
                    'n_catalogo_generale': form.n_catalogo_generale.data,
                    'n_catalogo_interno': form.n_catalogo_interno.data,
                    'n_catalogo_internazionale': form.n_catalogo_internazionale.data,
                    'soprintendenza': form.soprintendenza.data,

                    # TAB 2: Descrizioni
                    'd_stratigrafica': form.d_stratigrafica.data,
                    'd_interpretativa': form.d_interpretativa.data,
                    'descrizione': form.descrizione.data,
                    'interpretazione': form.interpretazione.data,
                    'osservazioni': form.osservazioni.data,

                    # TAB 3: Caratteristiche Fisiche
                    'formazione': form.formazione.data,
                    'stato_di_conservazione': form.stato_di_conservazione.data,
                    'colore': form.colore.data,
                    'consistenza': form.consistenza.data,
                    'struttura': form.struttura.data,
                    'quota_relativa': to_float(form.quota_relativa.data),
                    'quota_abs': to_float(form.quota_abs.data),
                    'lunghezza_max': to_float(form.lunghezza_max.data),
                    'larghezza_media': to_float(form.larghezza_media.data),
                    'altezza_max': to_float(form.altezza_max.data),
                    'altezza_min': to_float(form.altezza_min.data),
                    'profondita_max': to_float(form.profondita_max.data),
                    'profondita_min': to_float(form.profondita_min.data),

                    # TAB 4: Cronologia
                    'periodo_iniziale': form.periodo_iniziale.data,
                    'fase_iniziale': form.fase_iniziale.data,
                    'periodo_finale': form.periodo_finale.data,
                    'fase_finale': form.fase_finale.data,
                    'datazione': form.datazione.data,
                    'affidabilita': form.affidabilita.data,

                    # TAB 5: Relazioni Stratigrafiche
                    'rapporti': form.rapporti.data,

                    # TAB 6: Documentazione
                    'inclusi': form.inclusi.data,
                    'campioni': form.campioni.data,
                    'documentazione': form.documentazione.data,
                    'cont_per': form.cont_per.data,

                    # Altri campi
                    'flottazione': form.flottazione.data,
                    'setacciatura': form.setacciatura.data,
                }

                # Handle file upload for DOC units
                if form.unita_tipo.data == 'DOC' and form.documento_file.data:
                    file = form.documento_file.data
                    if file and file.filename:
                        import os
                        from werkzeug.utils import secure_filename
                        import datetime

                        # Create DoSC directory if it doesn't exist
                        dosc_dir = os.path.join(os.getcwd(), 'DoSC')
                        os.makedirs(dosc_dir, exist_ok=True)

                        # Generate filename: SITE_US_timestamp_originalname
                        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                        original_name = secure_filename(file.filename)
                        filename = f"{us_data['sito']}_{us_data['us']}_{timestamp}_{original_name}"

                        # Save file
                        file_path = os.path.join(dosc_dir, filename)
                        file.save(file_path)

                        # Store relative path in database
                        us_data['file_path'] = f"DoSC/{filename}"

                us = us_service.create_us(us_data)

                # Synchronize rapporti field to us_relationships_table
                try:
                    with db_manager.connection.get_session() as session:
                        relationship_sync_service.sync_rapporti_to_relationships_table(
                            sito=us_data['sito'],
                            us_number=int(us_data['us']),
                            rapporti_text=us_data.get('rapporti', ''),
                            session=session
                        )
                except Exception as sync_error:
                    print(f"Warning: Failed to sync relationships: {sync_error}")

                # Broadcast US creation (use data from form, not from detached instance)
                broadcast_us_created(socketio, us_data['sito'], us_data['us'])

                flash(f'US {us_data["us"]} creata con successo!', 'success')
                return redirect(url_for('us_list'))

            except Exception as e:
                flash(f'Errore nella creazione US: {str(e)}', 'error')
        elif request.method == 'POST':
            # Form validation failed - show errors
            flash('Errore nella validazione del form. Controlla i campi obbligatori.', 'error')
            print(f"Form validation errors: {form.errors}")
        
        return render_template('us/form.html', form=form, title='New SU')

    @app.route('/us/<us_id>/edit', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def edit_us(us_id):
        """Edit existing US - us_id is a composite string: sito__area__us_number"""
        from flask import session

        form = USForm()

        # Populate site choices
        sites = site_service.get_all_sites(size=100)
        form.sito.choices = [('', '-- Seleziona Sito --')] + [(s.sito, s.sito) for s in sites]

        # Populate datazione choices
        datazioni_choices = datazione_service.get_datazioni_choices()
        form.datazione.choices = [('', '-- Seleziona Datazione --')] + [(d['value'], d['label']) for d in datazioni_choices]

        # Populate thesaurus choices for US fields
        try:
            def_strat_values = thesaurus_service.get_field_values('us_table', 'definizione_stratigrafica')
            form.definizione_stratigrafica.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in def_strat_values]

            formazione_values = thesaurus_service.get_field_values('us_table', 'formazione')
            form.formazione.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in formazione_values]

            consistenza_values = thesaurus_service.get_field_values('us_table', 'consistenza')
            form.consistenza.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in consistenza_values]

            colore_values = thesaurus_service.get_field_values('us_table', 'colore')
            form.colore.choices = [('', '-- Seleziona --')] + [(v['value'], v['value']) for v in colore_values]
        except Exception as e:
            print(f"Error loading thesaurus values: {e}")
            # Fallback to empty choices
            form.definizione_stratigrafica.choices = [('', '-- Seleziona --')]
            form.formazione.choices = [('', '-- Seleziona --')]
            form.consistenza.choices = [('', '-- Seleziona --')]
            form.colore.choices = [('', '-- Seleziona --')]

        # Get existing US
        us = us_service.get_us_dto_by_id(us_id)
        if not us:
            flash('US non trovata', 'error')
            return redirect(url_for('us_list'))

        # Calculate prev/next navigation based on filters
        prev_id = None
        next_id = None
        current_position = 0
        total_records = 0

        try:
            # Get filters from session
            filters = session.get('us_filters', {})

            # Build query filters
            query_filters = {}
            if filters.get('sito'):
                query_filters['sito'] = filters['sito']
            if filters.get('area'):
                query_filters['area'] = filters['area']
            if filters.get('unita_tipo'):
                query_filters['unita_tipo'] = filters['unita_tipo']
            if filters.get('anno_scavo'):
                try:
                    query_filters['anno_scavo'] = int(filters['anno_scavo'])
                except (ValueError, TypeError):
                    pass
            if filters.get('us_number'):
                query_filters['us'] = filters['us_number']

            # Get full list of filtered US IDs
            all_us = us_service.get_all_us(size=10000, filters=query_filters)
            us_ids = [u.id_us for u in all_us]
            total_records = len(us_ids)

            # Find current position
            try:
                current_position = us_ids.index(us_id) + 1  # 1-based for display
                current_index = current_position - 1  # 0-based for array

                # Calculate prev/next
                if current_index > 0:
                    prev_id = us_ids[current_index - 1]
                if current_index < len(us_ids) - 1:
                    next_id = us_ids[current_index + 1]
            except ValueError:
                # Current US not in filtered list - this is ok, just no navigation
                pass
        except Exception as e:
            # If navigation calculation fails, just continue without it
            print(f"Navigation calculation error: {e}")
            pass

        if form.validate_on_submit():
            try:
                # Helper function to convert numeric fields
                def to_float(value):
                    if value and str(value).strip():
                        try:
                            return float(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                def to_int(value):
                    if value and str(value).strip():
                        try:
                            return int(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                update_data = {
                    # TAB 1: Informazioni Base
                    'sito': form.sito.data,
                    'area': form.area.data,
                    'us': to_int(form.us.data),
                    'unita_tipo': form.unita_tipo.data,
                    'tipo_documento': form.tipo_documento.data if form.unita_tipo.data == 'DOC' else None,
                    'anno_scavo': form.anno_scavo.data,
                    'scavato': form.scavato.data,
                    'schedatore': form.schedatore.data,
                    'metodo_di_scavo': form.metodo_di_scavo.data,
                    'data_schedatura': to_date(form.data_schedatura.data),
                    'attivita': form.attivita.data,
                    'direttore_us': form.direttore_us.data,
                    'responsabile_us': form.responsabile_us.data,
                    'settore': form.settore.data,
                    'quad_par': form.quad_par.data,
                    'ambient': form.ambient.data,
                    'saggio': form.saggio.data,
                    'n_catalogo_generale': form.n_catalogo_generale.data,
                    'n_catalogo_interno': form.n_catalogo_interno.data,
                    'n_catalogo_internazionale': form.n_catalogo_internazionale.data,
                    'soprintendenza': form.soprintendenza.data,

                    # TAB 2: Descrizioni
                    'd_stratigrafica': form.d_stratigrafica.data,
                    'd_interpretativa': form.d_interpretativa.data,
                    'descrizione': form.descrizione.data,
                    'interpretazione': form.interpretazione.data,
                    'osservazioni': form.osservazioni.data,

                    # TAB 3: Caratteristiche Fisiche
                    'formazione': form.formazione.data,
                    'stato_di_conservazione': form.stato_di_conservazione.data,
                    'colore': form.colore.data,
                    'consistenza': form.consistenza.data,
                    'struttura': form.struttura.data,
                    'quota_relativa': to_float(form.quota_relativa.data),
                    'quota_abs': to_float(form.quota_abs.data),
                    'lunghezza_max': to_float(form.lunghezza_max.data),
                    'larghezza_media': to_float(form.larghezza_media.data),
                    'altezza_max': to_float(form.altezza_max.data),
                    'altezza_min': to_float(form.altezza_min.data),
                    'profondita_max': to_float(form.profondita_max.data),
                    'profondita_min': to_float(form.profondita_min.data),

                    # TAB 4: Cronologia
                    'periodo_iniziale': form.periodo_iniziale.data,
                    'fase_iniziale': form.fase_iniziale.data,
                    'periodo_finale': form.periodo_finale.data,
                    'fase_finale': form.fase_finale.data,
                    'datazione': form.datazione.data,
                    'affidabilita': form.affidabilita.data,

                    # TAB 5: Relazioni Stratigrafiche
                    'rapporti': form.rapporti.data,

                    # TAB 6: Documentazione
                    'inclusi': form.inclusi.data,
                    'campioni': form.campioni.data,
                    'documentazione': form.documentazione.data,
                    'cont_per': form.cont_per.data,

                    # Altri campi
                    'flottazione': form.flottazione.data,
                    'setacciatura': form.setacciatura.data,
                }

                # Handle file upload for DOC units
                if form.unita_tipo.data == 'DOC' and form.documento_file.data:
                    file = form.documento_file.data
                    if file and file.filename:
                        import os
                        from werkzeug.utils import secure_filename
                        import datetime

                        # Create DoSC directory if it doesn't exist
                        dosc_dir = os.path.join(os.getcwd(), 'DoSC')
                        os.makedirs(dosc_dir, exist_ok=True)

                        # Generate filename: SITE_US_timestamp_originalname
                        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                        original_name = secure_filename(file.filename)
                        filename = f"{update_data['sito']}_{update_data['us']}_{timestamp}_{original_name}"

                        # Save file
                        file_path = os.path.join(dosc_dir, filename)
                        file.save(file_path)

                        # Store relative path in database
                        update_data['file_path'] = f"DoSC/{filename}"

                us_service.update_us(us_id, update_data)

                # Synchronize rapporti field to us_relationships_table
                try:
                    with db_manager.connection.get_session() as session:
                        relationship_sync_service.sync_rapporti_to_relationships_table(
                            sito=update_data['sito'],
                            us_number=int(update_data['us']),
                            rapporti_text=update_data.get('rapporti', ''),
                            session=session
                        )
                except Exception as sync_error:
                    print(f"Warning: Failed to sync relationships: {sync_error}")

                flash(f'US {update_data["us"]} aggiornata con successo!', 'success')
                return redirect(url_for('us_list'))

            except Exception as e:
                flash(f'Errore nell\'aggiornamento US: {str(e)}', 'error')

        # Pre-populate form with existing data
        elif request.method == 'GET':
            form.sito.data = us.sito
            form.area.data = us.area
            form.us.data = us.us
            form.unita_tipo.data = us.unita_tipo
            form.anno_scavo.data = us.anno_scavo
            form.scavato.data = us.scavato
            form.schedatore.data = us.schedatore
            form.metodo_di_scavo.data = us.metodo_di_scavo
            form.data_schedatura.data = us.data_schedatura
            form.attivita.data = us.attivita
            form.direttore_us.data = us.direttore_us
            form.responsabile_us.data = us.responsabile_us
            form.settore.data = us.settore
            form.quad_par.data = us.quad_par
            form.ambient.data = us.ambient
            form.saggio.data = us.saggio
            form.n_catalogo_generale.data = us.n_catalogo_generale
            form.n_catalogo_interno.data = us.n_catalogo_interno
            form.n_catalogo_internazionale.data = us.n_catalogo_internazionale
            form.soprintendenza.data = us.soprintendenza
            form.d_stratigrafica.data = us.d_stratigrafica
            form.d_interpretativa.data = us.d_interpretativa
            form.descrizione.data = us.descrizione
            form.interpretazione.data = us.interpretazione
            form.osservazioni.data = us.osservazioni
            form.formazione.data = us.formazione
            form.stato_di_conservazione.data = us.stato_di_conservazione
            form.colore.data = us.colore
            form.consistenza.data = us.consistenza
            form.struttura.data = us.struttura
            form.quota_relativa.data = us.quota_relativa
            form.quota_abs.data = us.quota_abs
            form.lunghezza_max.data = us.lunghezza_max
            form.larghezza_media.data = us.larghezza_media
            form.altezza_max.data = us.altezza_max
            form.altezza_min.data = us.altezza_min
            form.profondita_max.data = us.profondita_max
            form.profondita_min.data = us.profondita_min
            form.periodo_iniziale.data = us.periodo_iniziale
            form.fase_iniziale.data = us.fase_iniziale
            form.periodo_finale.data = us.periodo_finale
            form.fase_finale.data = us.fase_finale
            form.datazione.data = us.datazione
            form.affidabilita.data = us.affidabilita
            form.rapporti.data = us.rapporti
            form.inclusi.data = us.inclusi
            form.campioni.data = us.campioni
            form.documentazione.data = us.documentazione
            form.cont_per.data = us.cont_per
            form.flottazione.data = us.flottazione
            form.setacciatura.data = us.setacciatura

        # Get associated media files
        try:
            media_list = media_service.get_media_by_entity('us', us_id, size=100)
        except Exception as e:
            logger.warning(f"Failed to fetch media for US {us_id}: {e}")
            media_list = []

        return render_template('us/form.html',
                             form=form,
                             title='Edit SU',
                             edit_mode=True,
                             prev_id=prev_id,
                             next_id=next_id,
                             current_position=current_position,
                             total_records=total_records,
                             us_id=us_id,
                             media_list=media_list)

    # Inventory routes
    @app.route('/inventario')
    @login_required
    def inventario_list():
        page = request.args.get('page', 1, type=int)
        sito_filter = request.args.get('sito', '')
        tipo_filter = request.args.get('tipo', '')
        
        filters = {}
        if sito_filter:
            filters['sito'] = sito_filter
        if tipo_filter:
            filters['tipo_reperto'] = tipo_filter
        
        inventory_list = inventario_service.get_all_inventario(page=page, size=20, filters=filters)
        total = inventario_service.count_inventario(filters=filters)
        
        # Get options for filters
        sites = site_service.get_all_sites(size=100)
        
        return render_template('inventario/list.html', inventory_list=inventory_list,
                             sites=sites, total=total, page=page,
                             sito_filter=sito_filter, tipo_filter=tipo_filter)
    
    @app.route('/inventario/create', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def create_inventario():
        form = InventarioForm()

        # Populate site choices
        sites = site_service.get_all_sites(size=100)
        form.sito.choices = [('', '-- Seleziona Sito --')] + [(s.sito, s.sito) for s in sites]

        # Populate thesaurus choices
        form.tipo_reperto.choices = get_thesaurus_choices('tipo_reperto')
        form.stato_conservazione.choices = get_thesaurus_choices('stato_conservazione')
        form.corpo_ceramico.choices = get_thesaurus_choices('corpo_ceramico')
        form.rivestimento.choices = get_thesaurus_choices('rivestimento')

        if form.validate_on_submit():
            try:
                # Helper function to convert numeric fields
                def to_float(value):
                    if value and str(value).strip():
                        try:
                            return float(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                def to_int(value):
                    if value and str(value).strip():
                        try:
                            return int(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                inv_data = {
                    # TAB 1: Identificazione
                    'sito': form.sito.data,
                    'numero_inventario': form.numero_inventario.data,
                    'n_reperto': form.n_reperto.data,
                    'schedatore': form.schedatore.data,
                    'date_scheda': form.date_scheda.data,
                    'years': form.years.data,

                    # TAB 2: Classificazione
                    'tipo_reperto': form.tipo_reperto.data,
                    'criterio_schedatura': form.criterio_schedatura.data,
                    'definizione': form.definizione.data,
                    'tipo': form.tipo.data,
                    'tipo_contenitore': form.tipo_contenitore.data,
                    'struttura': form.struttura.data,
                    'descrizione': form.descrizione.data,

                    # TAB 3: Contesto
                    'area': form.area.data,
                    'us': form.us.data,
                    'punto_rinv': form.punto_rinv.data,
                    'elementi_reperto': form.elementi_reperto.data,

                    # TAB 4: Caratteristiche Fisiche
                    'stato_conservazione': form.stato_conservazione.data,
                    'lavato': form.lavato.data,
                    'nr_cassa': form.nr_cassa.data,
                    'luogo_conservazione': form.luogo_conservazione.data,

                    # TAB 5: Conservazione e Gestione
                    'repertato': form.repertato.data,
                    'diagnostico': form.diagnostico.data,

                    # TAB 6: Caratteristiche Ceramiche
                    'corpo_ceramico': form.corpo_ceramico.data,
                    'rivestimento': form.rivestimento.data,
                    'diametro_orlo': to_float(form.diametro_orlo.data),
                    'eve_orlo': to_float(form.eve_orlo.data),

                    # TAB 7: Misurazioni
                    'peso': to_float(form.peso.data),
                    'forme_minime': form.forme_minime.data,
                    'forme_massime': form.forme_massime.data,
                    'totale_frammenti': form.totale_frammenti.data,
                    'misurazioni': form.misurazioni.data,

                    # TAB 8: Documentazione
                    'datazione_reperto': form.datazione_reperto.data,
                    'rif_biblio': form.rif_biblio.data,
                    'tecnologie': form.tecnologie.data,
                    'negativo_photo': form.negativo_photo.data,
                    'diapositiva': form.diapositiva.data,
                }

                item = inventario_service.create_inventario(inv_data)

                # Broadcast inventario creation (use data from form, not from detached instance)
                broadcast_inventario_created(socketio, inv_data['numero_inventario'], inv_data['sito'])

                flash(f'Reperto {inv_data["numero_inventario"]} creato con successo!', 'success')
                return redirect(url_for('inventario_list'))

            except Exception as e:
                flash(f'Errore nella creazione reperto: {str(e)}', 'error')
        elif request.method == 'POST':
            # Form validation failed - show errors
            flash('Errore nella validazione del form. Controlla i campi obbligatori.', 'error')
            print(f"Form validation errors: {form.errors}")

        return render_template('inventario/form.html', form=form, title='New Find')

    @app.route('/inventario/<int:inv_id>/edit', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def edit_inventario(inv_id):
        """Edit existing inventario"""
        form = InventarioForm()

        # Populate site choices
        sites = site_service.get_all_sites(size=100)
        form.sito.choices = [('', '-- Seleziona Sito --')] + [(s.sito, s.sito) for s in sites]

        # Populate thesaurus choices
        form.tipo_reperto.choices = get_thesaurus_choices('tipo_reperto')
        form.stato_conservazione.choices = get_thesaurus_choices('stato_conservazione')
        form.corpo_ceramico.choices = get_thesaurus_choices('corpo_ceramico')
        form.rivestimento.choices = get_thesaurus_choices('rivestimento')

        # Get existing inventario using session to access all fields
        with db_manager.connection.get_session() as session:
            from pyarchinit_mini.models.inventario_materiali import InventarioMateriali as InvModel
            inv = session.query(InvModel).filter(InvModel.id_invmat == inv_id).first()

            if not inv:
                flash('Reperto non trovato', 'error')
                return redirect(url_for('inventario_list'))

            # Pre-populate form with existing data (inside session to avoid detached instance)
            if request.method == 'GET':
                form.sito.data = inv.sito
                form.numero_inventario.data = inv.numero_inventario
                form.n_reperto.data = inv.n_reperto
                form.schedatore.data = inv.schedatore
                form.date_scheda.data = inv.date_scheda
                form.years.data = inv.years
                form.tipo_reperto.data = inv.tipo_reperto
                form.criterio_schedatura.data = inv.criterio_schedatura
                form.definizione.data = inv.definizione
                form.tipo.data = inv.tipo
                form.tipo_contenitore.data = inv.tipo_contenitore
                form.struttura.data = inv.struttura
                form.descrizione.data = inv.descrizione
                form.area.data = inv.area
                form.us.data = inv.us
                form.punto_rinv.data = inv.punto_rinv
                form.elementi_reperto.data = inv.elementi_reperto
                form.stato_conservazione.data = inv.stato_conservazione
                form.lavato.data = inv.lavato
                form.nr_cassa.data = inv.nr_cassa
                form.luogo_conservazione.data = inv.luogo_conservazione
                form.repertato.data = inv.repertato
                form.diagnostico.data = inv.diagnostico
                form.corpo_ceramico.data = inv.corpo_ceramico
                form.rivestimento.data = inv.rivestimento
                form.diametro_orlo.data = inv.diametro_orlo
                form.eve_orlo.data = inv.eve_orlo
                form.peso.data = inv.peso
                form.forme_minime.data = inv.forme_minime
                form.forme_massime.data = inv.forme_massime
                form.totale_frammenti.data = inv.totale_frammenti
                form.misurazioni.data = inv.misurazioni
                form.datazione_reperto.data = inv.datazione_reperto
                form.rif_biblio.data = inv.rif_biblio
                form.tecnologie.data = inv.tecnologie
                form.negativo_photo.data = inv.negativo_photo
                form.diapositiva.data = inv.diapositiva

        if form.validate_on_submit():
            try:
                # Helper function to convert numeric fields
                def to_float(value):
                    if value and str(value).strip():
                        try:
                            return float(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                def to_int(value):
                    if value and str(value).strip():
                        try:
                            return int(value)
                        except (ValueError, TypeError):
                            return None
                    return None

                update_data = {
                    # TAB 1: Identificazione
                    'sito': form.sito.data,
                    'numero_inventario': form.numero_inventario.data,
                    'n_reperto': form.n_reperto.data,
                    'schedatore': form.schedatore.data,
                    'date_scheda': form.date_scheda.data,
                    'years': form.years.data,

                    # TAB 2: Classificazione
                    'tipo_reperto': form.tipo_reperto.data,
                    'criterio_schedatura': form.criterio_schedatura.data,
                    'definizione': form.definizione.data,
                    'tipo': form.tipo.data,
                    'tipo_contenitore': form.tipo_contenitore.data,
                    'struttura': form.struttura.data,
                    'descrizione': form.descrizione.data,

                    # TAB 3: Contesto
                    'area': form.area.data,
                    'us': form.us.data,
                    'punto_rinv': form.punto_rinv.data,
                    'elementi_reperto': form.elementi_reperto.data,

                    # TAB 4: Caratteristiche Fisiche
                    'stato_conservazione': form.stato_conservazione.data,
                    'lavato': form.lavato.data,
                    'nr_cassa': form.nr_cassa.data,
                    'luogo_conservazione': form.luogo_conservazione.data,

                    # TAB 5: Conservazione e Gestione
                    'repertato': form.repertato.data,
                    'diagnostico': form.diagnostico.data,

                    # TAB 6: Caratteristiche Ceramiche
                    'corpo_ceramico': form.corpo_ceramico.data,
                    'rivestimento': form.rivestimento.data,
                    'diametro_orlo': to_float(form.diametro_orlo.data),
                    'eve_orlo': to_float(form.eve_orlo.data),

                    # TAB 7: Misurazioni
                    'peso': to_float(form.peso.data),
                    'forme_minime': form.forme_minime.data,
                    'forme_massime': form.forme_massime.data,
                    'totale_frammenti': form.totale_frammenti.data,
                    'misurazioni': form.misurazioni.data,

                    # TAB 8: Documentazione
                    'datazione_reperto': form.datazione_reperto.data,
                    'rif_biblio': form.rif_biblio.data,
                    'tecnologie': form.tecnologie.data,
                    'negativo_photo': form.negativo_photo.data,
                    'diapositiva': form.diapositiva.data,
                }

                inventario_service.update_inventario(inv_id, update_data)

                flash(f'Reperto {update_data["numero_inventario"]} aggiornato con successo!', 'success')
                return redirect(url_for('inventario_list'))

            except Exception as e:
                flash(f'Errore nell\'aggiornamento reperto: {str(e)}', 'error')

        # Get associated media files
        try:
            media_list = media_service.get_media_by_entity('inventario', inv_id, size=100)
        except Exception as e:
            logger.warning(f"Failed to fetch media for inventario {inv_id}: {e}")
            media_list = []

        return render_template('inventario/form.html', form=form, title='Edit Find', edit_mode=True,
                             inv_id=inv_id, media_list=media_list)

    # Harris Matrix routes
    @app.route('/harris_matrix/<site_name>')
    def harris_matrix(site_name):
        """Harris Matrix - auto-selects best visualizer based on graph size"""
        try:
            print(f"🔵 [DEBUG] Starting harris_matrix for site: {site_name}")
            # Generate matrix
            graph = matrix_generator.generate_matrix(site_name)
            print(f"🔵 [DEBUG] Graph generated with {len(graph.nodes())} nodes")
            print(f"🔵 [DEBUG] Getting matrix levels...")
            levels = matrix_generator.get_matrix_levels(graph)
            print(f"🔵 [DEBUG] Getting matrix statistics...")
            stats = matrix_generator.get_matrix_statistics(graph)
            print(f"🔵 [DEBUG] Stats retrieved, preparing to render...")

            num_nodes = len(graph.nodes())
            print(f"🔵 [DEBUG] Checking graph size: {num_nodes} nodes")

            # For large graphs (> 500 nodes), direct rendering is too complex
            # Suggest using GraphML export instead (which works perfectly for large graphs)
            if num_nodes > 500:
                print(f"ℹ️  Large graph ({num_nodes} nodes) - too large for web rendering")
                print(f"   Please use GraphML export instead (works perfectly for large graphs)")

                # Show message to user to use GraphML export
                flash(f'This Harris Matrix has {num_nodes} nodes - too large for web rendering. ' +
                      'Please use GraphML export (available in the Export section) for best results with large graphs.',
                      'warning')

                return render_template('harris_matrix/large_graph_message.html',
                                     site_name=site_name,
                                     stats=stats,
                                     num_nodes=num_nodes)
            else:
                # Use matplotlib for small graphs
                matrix_image = matrix_visualizer.render_matplotlib(graph, levels)

                return render_template('harris_matrix/view.html',
                                     site_name=site_name,
                                     matrix_image=matrix_image,
                                     stats=stats,
                                     levels=levels,
                                     visualizer='matplotlib')

        except Exception as e:
            flash(f'Errore generazione Harris Matrix: {str(e)}', 'error')
            return redirect(url_for('sites_list'))

    @app.route('/harris_matrix/<site_name>/graphviz')
    def harris_matrix_graphviz(site_name):
        """Harris Matrix with Graphviz visualizer (desktop GUI style)"""
        try:
            # Get grouping parameter (period_area, period, area, none)
            grouping = request.args.get('grouping', 'period_area')

            # Generate matrix
            graph = matrix_generator.generate_matrix(site_name)
            levels = matrix_generator.get_matrix_levels(graph)
            stats = matrix_generator.get_matrix_statistics(graph)

            # Generate visualization with Graphviz
            output_path = graphviz_visualizer.create_matrix(
                graph,
                grouping=grouping,
                settings={
                    'show_legend': True,
                    'show_periods': grouping != 'none'
                }
            )

            # Read image and encode to base64
            with open(output_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')

            # Clean up temp file
            if os.path.exists(output_path):
                os.remove(output_path)

            return render_template('harris_matrix/view_graphviz.html',
                                 site_name=site_name,
                                 matrix_image=image_data,
                                 stats=stats,
                                 levels=levels,
                                 visualizer='graphviz',
                                 grouping=grouping)

        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'Errore generazione Harris Matrix Graphviz: {str(e)}', 'error')
            return redirect(url_for('sites_list'))

    # GraphML Export routes
    @app.route('/harris_matrix/graphml_export', methods=['GET', 'POST'])
    @login_required
    def export_harris_graphml():
        """Export Harris Matrix to GraphML format (yEd compatible)"""
        form = GraphMLExportForm()

        # Populate site choices
        sites = site_service.get_all_sites()
        form.site.choices = [(s.sito, s.sito) for s in sites]

        if form.validate_on_submit():
            try:
                site_name = form.site.data
                title = form.title.data or site_name
                grouping = form.grouping.data
                reverse_epochs = form.reverse_epochs.data

                # Use the new export_to_graphml method which:
                # 1. Generates Harris Matrix graph with transitive reduction
                # 2. Queries periodizzazione_table for datazione_estesa
                # 3. Creates DOT with cluster_datazione subgraphs
                # 4. Converts to GraphML with proper TableNode Rows

                # First generate the Harris Matrix graph
                graph = matrix_generator.generate_matrix(site_name)

                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.graphml', delete=False) as f:
                    temp_path = f.name

                result_path = matrix_generator.export_to_graphml(
                    graph=graph,
                    output_path=temp_path,
                    site_name=site_name,
                    title=title,
                    reverse_epochs=reverse_epochs
                )

                if not result_path:
                    flash('Errore durante l\'export GraphML', 'error')
                    return render_template('harris_matrix/graphml_export.html', form=form)

                # Send file
                filename = f"{site_name}_harris_matrix.graphml"
                response = send_file(
                    temp_path,
                    mimetype='application/xml',
                    as_attachment=True,
                    download_name=filename
                )

                # Clean up temp file after sending
                @response.call_on_close
                def cleanup():
                    try:
                        os.remove(temp_path)
                    except:
                        pass

                return response

            except Exception as e:
                import traceback
                traceback.print_exc()
                flash(f'Errore durante l\'export GraphML: {str(e)}', 'error')
                return render_template('harris_matrix/graphml_export.html', form=form)

        return render_template('harris_matrix/graphml_export.html', form=form)

    @app.route('/export/harris_matrix_filtered')
    @login_required
    def export_harris_matrix_filtered():
        """Export Harris Matrix GraphML with active filters from search"""
        from flask import session
        import networkx as nx

        try:
            # Get filters from session (set by us_list view)
            filters = session.get('us_filters', {})

            site_name = filters.get('sito')
            if not site_name:
                flash('Seleziona un sito per generare il Harris Matrix', 'warning')
                return redirect(url_for('us_list'))

            # Generate matrix for the site
            graph = matrix_generator.generate_matrix(site_name)

            # If filters are active (beyond site), filter the graph nodes
            if any([filters.get('area'), filters.get('unita_tipo'), filters.get('anno_scavo'), filters.get('us_number')]):
                # Get filtered US list
                with db_manager.connection.get_session() as db_session:
                    from pyarchinit_mini.models.us import US as USModel

                    query = db_session.query(USModel).filter(USModel.sito == site_name)

                    if filters.get('area'):
                        query = query.filter(USModel.area == filters['area'])
                    if filters.get('unita_tipo'):
                        query = query.filter(USModel.unita_tipo == filters['unita_tipo'])
                    if filters.get('anno_scavo'):
                        try:
                            query = query.filter(USModel.anno_scavo == int(filters['anno_scavo']))
                        except (ValueError, TypeError):
                            pass
                    if filters.get('us_number'):
                        query = query.filter(USModel.us == filters['us_number'])

                    filtered_us = query.all()
                    filtered_us_numbers = set([str(us.us) for us in filtered_us])

                # Create subgraph with only filtered nodes and their relationships
                filtered_graph = nx.DiGraph()

                # Add filtered nodes with their attributes
                for node in graph.nodes():
                    node_us = str(node)
                    if node_us in filtered_us_numbers:
                        filtered_graph.add_node(node, **graph.nodes[node])

                # Add edges only between filtered nodes
                for u, v in graph.edges():
                    if str(u) in filtered_us_numbers and str(v) in filtered_us_numbers:
                        filtered_graph.add_edge(u, v, **graph.edges[u, v])

                graph = filtered_graph

            if len(graph.nodes()) == 0:
                flash('Nessuna US trovata con i filtri attuali per generare il Harris Matrix', 'warning')
                return redirect(url_for('us_list'))

            # Export to GraphML
            import tempfile
            with tempfile.NamedTemporaryFile(mode='w', suffix='.graphml', delete=False) as f:
                temp_path = f.name

            # Generate descriptive title based on filters
            title_parts = [site_name]
            if filters.get('area'):
                title_parts.append(f"Area {filters['area']}")
            if filters.get('unita_tipo'):
                title_parts.append(filters['unita_tipo'])
            if filters.get('anno_scavo'):
                title_parts.append(str(filters['anno_scavo']))
            title = ' - '.join(title_parts) + ' (Filtered)'

            result_path = matrix_generator.export_to_graphml(
                graph=graph,
                output_path=temp_path,
                site_name=site_name,
                title=title,
                reverse_epochs=False
            )

            if not result_path:
                flash('Errore durante l\'export GraphML', 'error')
                return redirect(url_for('us_list'))

            # Generate filename with filter indicators
            filename_parts = [site_name]
            if filters.get('area'):
                filename_parts.append(f"area_{filters['area']}")
            if filters.get('unita_tipo'):
                filename_parts.append(filters['unita_tipo'])
            if filters.get('anno_scavo'):
                filename_parts.append(str(filters['anno_scavo']))
            filename = '_'.join(filename_parts) + '_harris_matrix_filtered.graphml'

            # Send file
            response = send_file(
                temp_path,
                mimetype='application/xml',
                as_attachment=True,
                download_name=filename
            )

            # Clean up temp file after sending
            @response.call_on_close
            def cleanup():
                try:
                    os.remove(temp_path)
                except:
                    pass

            return response

        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'Errore durante l\'export GraphML: {str(e)}', 'error')
            return redirect(url_for('us_list'))

    # Stratigraphic Validation routes
    @app.route('/validate/<site_name>')
    def validate_stratigraphic(site_name):
        """Validate stratigraphic relationships for a site"""
        try:
            # Get all US for the site
            filters = {'sito': site_name}
            us_list = us_service.get_all_us(size=1000, filters=filters)

            # Initialize validator
            validator = StratigraphicValidator()

            # Add all units to validator
            for us in us_list:
                us_num = getattr(us, 'us', None)
                if us_num:
                    validator.add_unit(us_num, {
                        'sito': getattr(us, 'sito', ''),
                        'area': getattr(us, 'area', ''),
                        'd_stratigrafica': getattr(us, 'd_stratigrafica', ''),
                        'rapporti': getattr(us, 'rapporti', '')
                    })

                    # Parse and add relationships
                    rapporti = getattr(us, 'rapporti', '')
                    if rapporti:
                        relationships = validator.parse_relationships(us_num, rapporti)
                        for rel_type, target_us in relationships:
                            validator.add_relationship(us_num, target_us, rel_type)

            # Convert US list to dict format for validation
            us_list_dicts = []
            us_map = {}  # Map to store US id by (site, area, us_number)
            for us in us_list:
                us_num = getattr(us, 'us', None)
                us_dict = {
                    'id_us': getattr(us, 'id_us', None),
                    'sito': getattr(us, 'sito', ''),
                    'area': getattr(us, 'area', ''),
                    'us': us_num,
                    'rapporti': getattr(us, 'rapporti', ''),
                    'd_interpretativa': getattr(us, 'd_interpretativa', '')
                }
                us_list_dicts.append(us_dict)
                if us_num:
                    us_map[(us_dict['sito'], us_dict['area'], us_num)] = getattr(us, 'id_us', None)

            # Load periodization data for chronological validation
            periodization_list = []
            try:
                from pyarchinit_mini.models.harris_matrix import Periodizzazione, Period
                with db_manager.connection.get_session() as session:
                    # Query periodization with period dates
                    period_data = session.query(
                        Periodizzazione,
                        Period.start_date.label('period_start'),
                        Period.end_date.label('period_end')
                    ).outerjoin(
                        Period,
                        Periodizzazione.period_id_final == Period.id_period
                    ).filter(
                        Periodizzazione.sito == site_name
                    ).all()

                    for perio, period_start, period_end in period_data:
                        # Fallback to initial period if final not available
                        if period_end is None and perio.period_id_initial:
                            initial_period = session.query(Period).filter(
                                Period.id_period == perio.period_id_initial
                            ).first()
                            if initial_period:
                                period_start = initial_period.start_date
                                period_end = initial_period.end_date

                        periodization_list.append({
                            'sito': perio.sito,
                            'area': perio.area or '',
                            'us': perio.us,
                            'periodo_iniziale': perio.periodo_iniziale,
                            'fase_iniziale': perio.fase_iniziale,
                            'periodo_finale': perio.periodo_finale,
                            'fase_finale': perio.fase_finale,
                            'datazione_estesa': perio.datazione_estesa,
                            'start_date': period_start,
                            'end_date': period_end
                        })
            except Exception as e:
                # If periodization loading fails, continue without chronological validation
                print(f"Warning: Could not load periodization data: {e}")
                periodization_list = []

            # Get validation report with chronological validation
            report = validator.get_validation_report(us_list_dicts, periodization_list)

            # Extract cycles from errors
            cycles = []
            non_cycle_errors = []
            for error in report.get('errors', []):
                if 'cycle' in error.lower() or 'ciclo' in error.lower():
                    # Extract cycle from error message
                    # Format: "Temporal paradox (cycle): 1 → 2 → 3 → 1"
                    if '→' in error:
                        cycle_part = error.split(':')[-1].strip()
                        cycle_us = [us.strip() for us in cycle_part.split('→')]
                        cycles.append(cycle_us[:-1])  # Remove duplicate last element
                else:
                    non_cycle_errors.append(error)

            # Generate relationship fixes to find missing reciprocals
            fixes = validator.generate_relationship_fixes(us_list_dicts)

            # Convert fixes to display format
            missing_reciprocals = []
            for update in fixes.get('updates', []):
                # Parse the update to extract relationship info
                from_us = update['us']
                reason = update.get('reason', '')

                # Extract target US and relationship from reason
                # Format: "Aggiunta relazione reciproca: US X copre Y"
                if 'US' in reason:
                    parts = reason.split(':')[-1].strip().split()
                    if len(parts) >= 3:
                        try:
                            source_us = parts[1]  # US number after "US"
                            rel_type = ' '.join(parts[2:-1])  # Relationship type
                            target_us = parts[-1]  # Target US

                            # Determine expected reciprocal
                            inverse_rel = validator.INVERSE_RELATIONS.get(rel_type, '')

                            missing_reciprocals.append({
                                'from_us': source_us,
                                'to_us': from_us,
                                'relationship': rel_type,
                                'expected_reciprocal': inverse_rel
                            })
                        except (IndexError, ValueError):
                            pass

            # Get statistics
            stats = {
                'total_us': len(us_list),
                'total_relationships': report.get('relationships_found', 0),
                'total_errors': len(non_cycle_errors),
                'total_cycles': len(cycles),
                'missing_reciprocals': len(missing_reciprocals),
                'is_valid': report.get('valid', False) and len(cycles) == 0 and len(missing_reciprocals) == 0
            }

            # Return complete report
            return render_template('validation/report.html',
                                 site_name=site_name,
                                 stats=stats,
                                 errors=non_cycle_errors,
                                 cycles=cycles,
                                 missing_reciprocals=missing_reciprocals)

        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'Errore validazione: {str(e)}', 'error')
            return redirect(url_for('sites_list'))

    @app.route('/validate/<site_name>/fix', methods=['POST'])
    def fix_stratigraphic(site_name):
        """Apply automatic fixes to stratigraphic relationships"""
        try:
            fix_type = request.form.get('fix_type', 'reciprocals')

            # Get all US for the site
            filters = {'sito': site_name}
            us_list_objects = us_service.get_all_us(size=1000, filters=filters)

            # Convert to dict for validator
            us_list = []
            for us in us_list_objects:
                us_dict = {
                    'sito': getattr(us, 'sito', ''),
                    'area': getattr(us, 'area', ''),
                    'us': getattr(us, 'us', None),
                    'rapporti': getattr(us, 'rapporti', '')
                }
                us_list.append(us_dict)

            # Initialize validator
            validator = StratigraphicValidator()

            if fix_type == 'reciprocals':
                # Generate fixes for missing reciprocals
                fixes = validator.generate_relationship_fixes(us_list)

                # Apply updates to existing US
                updates_count = 0
                for update in fixes.get('updates', []):
                    us_num = update['us']
                    new_rapporti = update['new_value']  # Fixed: was update['rapporti']

                    # Update US in database
                    # Find the US by site and us number
                    us_to_update = None
                    for us_obj in us_list:
                        if getattr(us_obj, 'us', None) == us_num:
                            us_to_update = us_obj
                            break

                    if us_to_update:
                        us_service.update_us_by_id(
                            getattr(us_to_update, 'id_us'),
                            {'rapporti': new_rapporti}
                        )
                        updates_count += 1

                # Create new US if needed
                creates_count = len(fixes.get('creates', []))

                flash(f'Fix applicati: {updates_count} US aggiornate, {creates_count} US da creare manualmente', 'success')

            return redirect(url_for('validate_stratigraphic', site_name=site_name))

        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'Errore applicazione fix: {str(e)}', 'error')
            return redirect(url_for('validate_stratigraphic', site_name=site_name))

    # Export routes
    @app.route('/export/site_pdf/<int:site_id>')
    def export_site_pdf(site_id):
        try:
            # Get site data and convert to dict within session scope
            with db_manager.connection.get_session() as session:
                from pyarchinit_mini.models.site import Site as SiteModel
                from pyarchinit_mini.models.us import US as USModel
                from pyarchinit_mini.models.inventario_materiali import InventarioMateriali as InvModel

                site = session.query(SiteModel).filter(SiteModel.id_sito == site_id).first()
                if not site:
                    flash('Sito non trovato', 'error')
                    return redirect(url_for('sites_list'))

                site_name = site.sito
                site_dict = site.to_dict()

                # Get related data and convert to dict within session
                us_records = session.query(USModel).filter(USModel.sito == site_name).limit(100).all()
                us_list = [us.to_dict() for us in us_records]

                inv_records = session.query(InvModel).filter(InvModel.sito == site_name).limit(100).all()
                inventory_list = [inv.to_dict() for inv in inv_records]

            # Generate PDF (outside session)
            pdf_bytes = pdf_generator.generate_site_report(
                site_dict,
                us_list,
                inventory_list
            )

            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(pdf_bytes)
                tmp_path = tmp.name

            return send_file(tmp_path, as_attachment=True,
                           download_name=f"relazione_{site_name}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF: {str(e)}', 'error')
            return redirect(url_for('view_site', site_id=site_id))

    @app.route('/export/site_pdf_with_matrix/<site_name>')
    def export_site_pdf_with_matrix(site_name):
        """Export site PDF with integrated Harris Matrix"""
        try:
            # Generate Harris Matrix image
            graph = matrix_generator.generate_matrix(site_name)
            stats = matrix_generator.get_matrix_statistics(graph)

            # Create temporary file for matrix image
            matrix_img_path = graphviz_visualizer.create_matrix(
                graph,
                grouping='period_area',
                settings={'show_legend': True, 'show_periods': True}
            )

            # Generate PDF with matrix
            pdf_bytes = pdf_generator.generate_harris_matrix_report(
                site_name,
                matrix_img_path,
                stats
            )

            # Cleanup temp image
            if os.path.exists(matrix_img_path):
                os.remove(matrix_img_path)

            # Create temporary PDF file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(pdf_bytes)
                tmp_path = tmp.name

            return send_file(tmp_path, as_attachment=True,
                           download_name=f"harris_matrix_{site_name}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF Harris Matrix: {str(e)}', 'error')
            return redirect(url_for('harris_matrix', site_name=site_name))

    @app.route('/export/us_pdf')
    def export_us_pdf():
        """Export US list PDF with active filters"""
        from flask import session

        try:
            # Get filters from session (set by us_list view)
            filters = session.get('us_filters', {})

            # Build query filters
            query_filters = {}
            if filters.get('sito'):
                query_filters['sito'] = filters['sito']
            if filters.get('area'):
                query_filters['area'] = filters['area']
            if filters.get('unita_tipo'):
                query_filters['unita_tipo'] = filters['unita_tipo']
            if filters.get('anno_scavo'):
                try:
                    query_filters['anno_scavo'] = int(filters['anno_scavo'])
                except (ValueError, TypeError):
                    pass
            if filters.get('us_number'):
                query_filters['us'] = filters['us_number']

            # Get US data within session
            with db_manager.connection.get_session() as db_session:
                from pyarchinit_mini.models.us import US as USModel

                query = db_session.query(USModel)

                # Apply all filters
                if query_filters.get('sito'):
                    query = query.filter(USModel.sito == query_filters['sito'])
                if query_filters.get('area'):
                    query = query.filter(USModel.area == query_filters['area'])
                if query_filters.get('unita_tipo'):
                    query = query.filter(USModel.unita_tipo == query_filters['unita_tipo'])
                if query_filters.get('anno_scavo'):
                    query = query.filter(USModel.anno_scavo == query_filters['anno_scavo'])
                if query_filters.get('us'):
                    query = query.filter(USModel.us == query_filters['us'])

                us_records = query.limit(500).all()
                us_list = [us.to_dict() for us in us_records]

            if not us_list:
                flash('Nessuna US trovata con i filtri attuali', 'warning')
                return redirect(url_for('us_list'))

            # Generate PDF filename based on filters
            site_name_clean = filters.get('sito', 'Filtered_Results')
            filter_suffix = ''
            if filters.get('area'):
                filter_suffix += f'_area_{filters["area"]}'
            if filters.get('unita_tipo'):
                filter_suffix += f'_{filters["unita_tipo"]}'
            if filters.get('anno_scavo'):
                filter_suffix += f'_{filters["anno_scavo"]}'

            # Create temporary file for PDF output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                output_path = tmp.name

            # Generate PDF to file
            pdf_generator.generate_us_pdf(site_name_clean, us_list, output_path)

            return send_file(output_path, as_attachment=True,
                           download_name=f"us_{site_name_clean}{filter_suffix}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF US: {str(e)}', 'error')
            return redirect(url_for('us_list'))

    @app.route('/export/us_single_pdf/<sito>/<int:us_number>')
    def export_us_single_pdf(sito, us_number):
        """Export single US PDF"""
        try:
            # Get US data within session
            with db_manager.connection.get_session() as session:
                from pyarchinit_mini.models.us import US as USModel

                us = session.query(USModel).filter(
                    USModel.sito == sito,
                    USModel.us == us_number
                ).first()

                if not us:
                    flash('US non trovata', 'error')
                    return redirect(url_for('us_list'))

                us_dict = us.to_dict()

            # Create temporary file for PDF output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                output_path = tmp.name

            # Generate PDF with single US
            pdf_generator.generate_us_pdf(sito, [us_dict], output_path)

            return send_file(output_path, as_attachment=True,
                           download_name=f"us_{sito}_{us_number}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF US: {str(e)}', 'error')
            return redirect(url_for('us_list'))

    @app.route('/export/inventario_pdf')
    def export_inventario_pdf():
        """Export Inventario list PDF"""
        try:
            site_name = request.args.get('sito')

            # Get inventario data within session
            with db_manager.connection.get_session() as session:
                from pyarchinit_mini.models.inventario_materiali import InventarioMateriali as InvModel

                query = session.query(InvModel)
                if site_name:
                    query = query.filter(InvModel.sito == site_name)

                inv_records = query.limit(500).all()
                inventory_list = [inv.to_dict() for inv in inv_records]

            if not inventory_list:
                flash('Nessun reperto trovato', 'warning')
                return redirect(url_for('inventario_list'))

            # Generate PDF
            site_name_clean = site_name if site_name else 'Tutti_i_siti'

            # Create temporary file for PDF output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                output_path = tmp.name

            # Generate PDF to file
            pdf_generator.generate_inventario_pdf(site_name_clean, inventory_list, output_path)

            return send_file(output_path, as_attachment=True,
                           download_name=f"inventario_{site_name_clean}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF Inventario: {str(e)}', 'error')
            return redirect(url_for('inventario_list'))

    @app.route('/export/inventario_single_pdf/<int:inv_id>')
    def export_inventario_single_pdf(inv_id):
        """Export single Inventario PDF"""
        try:
            # Get inventario data within session
            with db_manager.connection.get_session() as session:
                from pyarchinit_mini.models.inventario_materiali import InventarioMateriali as InvModel

                inv = session.query(InvModel).filter(InvModel.id_invmat == inv_id).first()

                if not inv:
                    flash('Reperto non trovato', 'error')
                    return redirect(url_for('inventario_list'))

                inv_dict = inv.to_dict()
                site_name = inv_dict.get('sito', 'Unknown')

            # Create temporary file for PDF output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                output_path = tmp.name

            # Generate PDF with single item
            pdf_generator.generate_inventario_pdf(site_name, [inv_dict], output_path)

            return send_file(output_path, as_attachment=True,
                           download_name=f"inventario_{inv_id}.pdf",
                           mimetype='application/pdf')

        except Exception as e:
            flash(f'Errore export PDF Inventario: {str(e)}', 'error')
            return redirect(url_for('inventario_list'))

    # Media routes
    # API endpoints for media upload entity selection
    @app.route('/api/media/sites')
    @login_required
    def api_media_sites():
        """Get list of all sites for media upload form"""
        try:
            sites = site_service.get_all_sites(size=1000)
            return jsonify([{
                'id': site.id_sito,
                'name': site.sito
            } for site in sites])
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/media/us')
    @login_required
    def api_media_us():
        """Get list of all US for media upload form"""
        try:
            site_name = request.args.get('site')
            filters = {}
            if site_name:
                filters['sito'] = site_name

            us_list = us_service.get_all_us(size=1000, filters=filters if filters else None)

            return jsonify([{
                'id': us.id_us,
                'sito': us.sito,
                'area': us.area or '',
                'us': str(us.us),
                'label': f"{us.sito} - Area {us.area or 'N/A'} - US {us.us}"
            } for us in us_list])
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/media/inventario')
    @login_required
    def api_media_inventario():
        """Get list of all inventory items for media upload form"""
        try:
            site_name = request.args.get('site')
            filters = {}
            if site_name:
                filters['sito'] = site_name

            inv_list = inventario_service.get_all_inventario(size=1000, filters=filters if filters else None)

            return jsonify([{
                'id': inv.id_invmat,
                'sito': inv.sito,
                'numero_inventario': inv.numero_inventario,
                'label': f"{inv.sito} - Inv. {inv.numero_inventario}"
            } for inv in inv_list])
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/media/upload', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def upload_media():
        form = MediaUploadForm()

        # Populate site dropdown choices
        sites = site_service.get_all_sites(size=1000)
        site_choices = [('', '-- Seleziona Sito --')] + [(str(site.id_sito), site.sito) for site in sites]
        form.site_id.choices = [(int(id_val) if id_val else 0, name) for id_val, name in site_choices if id_val]
        form.us_site.choices = [('', '-- Seleziona Sito --')] + [(site.sito, site.sito) for site in sites]
        form.inv_site.choices = [('', '-- Seleziona Sito --')] + [(site.sito, site.sito) for site in sites]

        if form.validate_on_submit():
            try:
                # Determine entity_id based on entity_type and form fields
                entity_type = form.entity_type.data
                entity_id = None

                if entity_type == 'site':
                    entity_id = form.site_id.data
                elif entity_type == 'us':
                    # Find US by site, area, and US number
                    site_name = form.us_site.data
                    area = form.us_area.data
                    us_number = form.us_number.data

                    if not all([site_name, us_number]):
                        flash('For US, Site and US Number are required', 'error')
                        return render_template('media/upload.html', form=form)

                    # Search for US record
                    us_filters = {'sito': site_name, 'us': us_number}
                    if area:
                        us_filters['area'] = area
                    us_list = us_service.get_all_us(size=1000, filters=us_filters)

                    if not us_list:
                        flash(f'US not found: {site_name} - Area {area or "N/A"} - US {us_number}', 'error')
                        return render_template('media/upload.html', form=form)

                    entity_id = us_list[0].id_us

                elif entity_type == 'inventario':
                    # Find inventory item by site and inventory number
                    site_name = form.inv_site.data
                    inv_number = form.inv_number.data

                    if not all([site_name, inv_number]):
                        flash('For Inventory, Site and Inventory Number are required', 'error')
                        return render_template('media/upload.html', form=form)

                    # Search for inventory record
                    inv_list = inventario_service.get_all_inventario(size=1000, filters={'sito': site_name, 'numero_inventario': inv_number})

                    if not inv_list:
                        flash(f'Inventory not found: {site_name} - Inv. {inv_number}', 'error')
                        return render_template('media/upload.html', form=form)

                    entity_id = inv_list[0].id_invmat

                if not entity_id:
                    flash('Error: unable to determine associated entity', 'error')
                    return render_template('media/upload.html', form=form)

                uploaded_file = form.file.data
                if uploaded_file and uploaded_file.filename:
                    # Save uploaded file temporarily
                    filename = secure_filename(uploaded_file.filename)
                    temp_path = os.path.join(tempfile.gettempdir(), filename)
                    uploaded_file.save(temp_path)

                    # Store using media handler
                    metadata = media_handler.store_file(
                        temp_path,
                        entity_type,
                        entity_id,
                        form.description.data,
                        "",  # tags
                        form.author.data
                    )

                    # Remove thumbnail_path - thumbnails go in MediaThumb table, not Media table
                    metadata.pop('thumbnail_path', None)

                    # Save record to database
                    media_record = media_service.create_media_record(metadata)

                    # Get ID immediately while object is in memory
                    media_id = media_record.id_media

                    # Clean up temp file
                    os.remove(temp_path)

                    flash(f'File uploaded successfully! (ID: {media_id})', 'success')
                    return redirect(url_for('media_list'))

            except Exception as e:
                flash(f'File upload error: {str(e)}', 'error')

        # Pre-populate form from query parameters (when coming from entity forms)
        if request.method == 'GET':
            entity_type_param = request.args.get('entity_type')
            entity_id_param = request.args.get('entity_id')

            if entity_type_param:
                form.entity_type.data = entity_type_param

                if entity_id_param and entity_type_param == 'site':
                    try:
                        form.site_id.data = int(entity_id_param)
                    except (ValueError, TypeError):
                        pass

        return render_template('media/upload.html', form=form)

    @app.route('/media/upload/ajax', methods=['POST'])
    @login_required
    @write_permission_required
    def upload_media_ajax():
        """AJAX endpoint for drag-and-drop media uploads"""
        try:
            entity_type = request.form.get('entity_type')
            entity_id = request.form.get('entity_id')

            if not entity_type or not entity_id:
                return jsonify({'error': f'Missing entity_type ({entity_type}) or entity_id ({entity_id})'}), 400

            # Convert entity_id to int
            try:
                entity_id = int(entity_id)
            except ValueError:
                return jsonify({'error': f'Invalid entity_id: {entity_id}'}), 400

            files = request.files.getlist('files')
            if not files or len(files) == 0:
                return jsonify({'error': 'No files provided'}), 400

            uploaded_ids = []

            for uploaded_file in files:
                if uploaded_file and uploaded_file.filename:
                    # Save uploaded file temporarily
                    filename = secure_filename(uploaded_file.filename)
                    temp_path = os.path.join(tempfile.gettempdir(), filename)
                    uploaded_file.save(temp_path)

                    # Store using media handler
                    metadata = media_handler.store_file(
                        temp_path,
                        entity_type,
                        entity_id,
                        "",  # description (empty for drag-drop)
                        "",  # tags
                        ""   # author
                    )

                    # Remove thumbnail_path
                    metadata.pop('thumbnail_path', None)

                    # Save record to database
                    media_record = media_service.create_media_record(metadata)
                    uploaded_ids.append(media_record.id_media)

                    # Clean up temp file
                    os.remove(temp_path)

            return jsonify({
                'success': True,
                'uploaded': len(uploaded_ids),
                'ids': uploaded_ids
            }), 200

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/media/list')
    @login_required
    def media_list():
        """List all uploaded media files with search and filter capabilities"""
        try:
            # Get query parameters
            search_term = request.args.get('search', '').strip()
            media_type_filter = request.args.get('type', '').strip()
            entity_type_filter = request.args.get('entity', '').strip()
            page = int(request.args.get('page', 1))
            per_page = int(request.args.get('per_page', 20))

            # Get media list based on filters
            if search_term:
                media_list = media_service.search_media(search_term, page=page, size=per_page)
            elif media_type_filter:
                media_list = media_service.get_media_by_type(media_type_filter, page=page, size=per_page)
            elif entity_type_filter:
                # For entity filter, we need to get all media and filter (or modify service)
                media_list = media_service.get_all_media(page=page, size=per_page,
                                                        filters={'entity_type': entity_type_filter} if entity_type_filter else None)
            else:
                media_list = media_service.get_all_media(page=page, size=per_page)

            # Get statistics
            stats = media_service.get_media_statistics()

            # Enrich media with entity information
            # First, capture entity info while objects are in memory
            media_enrichment = []
            for media in media_list:
                try:
                    entity_type = media.entity_type
                    entity_id = media.entity_id
                    media_enrichment.append((media, entity_type, entity_id))
                except:
                    media_enrichment.append((media, None, None))

            # Now enrich with entity names
            for media, entity_type, entity_id in media_enrichment:
                try:
                    if entity_type == 'site':
                        site = site_service.get_site_by_id(entity_id)
                        media.entity_name = site.sito if site else f'Site ID {entity_id}'
                    elif entity_type == 'us':
                        us = us_service.get_us_by_id(entity_id)
                        media.entity_name = f"{us.sito} - US {us.us}" if us else f'US ID {entity_id}'
                    elif entity_type == 'inventario':
                        inv = inventario_service.get_inventario_by_id(entity_id)
                        media.entity_name = f"{inv.sito} - {inv.numero_inventario}" if inv else f'Inventory ID {entity_id}'
                    else:
                        media.entity_name = f'{entity_type} ID {entity_id}' if entity_type else 'Unknown'
                except:
                    media.entity_name = f'{entity_type} ID {entity_id}' if entity_type else 'Unknown'

            return render_template('media/list.html',
                                 media_list=media_list,
                                 stats=stats,
                                 search_term=search_term,
                                 media_type_filter=media_type_filter,
                                 entity_type_filter=entity_type_filter,
                                 page=page,
                                 per_page=per_page)
        except Exception as e:
            flash(f'Error loading media list: {str(e)}', 'error')
            return render_template('media/list.html', media_list=[], stats={})

    @app.route('/media/<path:filepath>')
    @login_required
    def serve_media(filepath):
        """Serve media files from the media directory"""
        from flask import send_from_directory
        import os

        # Get the media directory path from centralized location
        from pathlib import Path
        media_dir = str(Path.home() / '.pyarchinit_mini' / 'media')
        return send_from_directory(media_dir, filepath)

    @app.route('/media/delete/<int:media_id>', methods=['POST'])
    @login_required
    @write_permission_required
    def delete_media(media_id):
        """Delete a media file"""
        try:
            # Get the media record to retrieve entity info before deletion
            media = media_service.get_media_by_id(media_id)
            if not media:
                flash('Media file not found', 'error')
                return redirect(url_for('media_list'))

            # Store entity info for potential redirect
            entity_type = media.entity_type
            entity_id = media.entity_id

            # Delete the media
            success = media_service.delete_media(media_id, delete_file=True)

            if success:
                flash('Media file deleted successfully', 'success')
            else:
                flash('Error deleting media file', 'error')

            # Redirect back to referrer or media list
            referrer = request.referrer
            if referrer and ('edit_site' in referrer or 'edit_us' in referrer or 'edit_inventario' in referrer):
                return redirect(referrer)
            else:
                return redirect(url_for('media_list'))

        except Exception as e:
            flash(f'Error deleting media: {str(e)}', 'error')
            return redirect(url_for('media_list'))

    @app.route('/media/view/spreadsheet/<int:media_id>')
    @login_required
    def view_spreadsheet(media_id):
        """View Excel/CSV files as HTML table"""
        try:
            import pandas as pd
            from pathlib import Path

            # Get the media record
            media = media_service.get_media_by_id(media_id)
            if not media:
                return '<html><body><h3>Media file not found</h3></body></html>', 404

            # Get file path
            if not media.media_path:
                return '<html><body><h3>File path not available</h3></body></html>', 404

            file_path = Path(media.media_path)
            if not file_path.exists():
                return '<html><body><h3>File not found on disk</h3></body></html>', 404

            # Read file based on extension
            file_ext = file_path.suffix.lower()

            try:
                if file_ext in ['.xlsx', '.xls']:
                    # Read Excel file
                    df = pd.read_excel(file_path)
                elif file_ext == '.csv':
                    # Read CSV file
                    df = pd.read_csv(file_path)
                else:
                    return f'<html><body><h3>Unsupported file format: {file_ext}</h3></body></html>', 400

                # Convert to HTML with Bootstrap styling
                html_table = df.to_html(
                    classes='table table-striped table-bordered table-hover table-sm',
                    index=False,
                    escape=True,
                    max_rows=1000  # Limit to 1000 rows for performance
                )

                # Create complete HTML page
                html_content = f'''
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <title>{media.media_name}</title>
                    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
                    <style>
                        body {{
                            padding: 20px;
                            font-size: 12px;
                        }}
                        .table-container {{
                            overflow-x: auto;
                            max-width: 100%;
                        }}
                        .table {{
                            margin-bottom: 0;
                        }}
                        .table thead th {{
                            position: sticky;
                            top: 0;
                            background-color: #212529;
                            color: white;
                            z-index: 10;
                        }}
                        h3 {{
                            margin-bottom: 20px;
                        }}
                        .info {{
                            margin-bottom: 15px;
                            padding: 10px;
                            background-color: #f8f9fa;
                            border-left: 4px solid #0d6efd;
                        }}
                    </style>
                </head>
                <body>
                    <div class="container-fluid">
                        <h3>{media.media_name}</h3>
                        <div class="info">
                            <p class="mb-1"><strong>Rows:</strong> {len(df)} | <strong>Columns:</strong> {len(df.columns)}</p>
                            {f'<p class="mb-0 text-muted"><small>Showing first 1000 rows</small></p>' if len(df) > 1000 else ''}
                        </div>
                        <div class="table-container">
                            {html_table}
                        </div>
                    </div>
                </body>
                </html>
                '''

                return html_content

            except Exception as e:
                return f'<html><body><h3>Error reading file: {str(e)}</h3></body></html>', 500

        except Exception as e:
            return f'<html><body><h3>Error: {str(e)}</h3></body></html>', 500

    @app.route('/media/view/docx/<int:media_id>')
    @login_required
    def view_docx(media_id):
        """View DOCX files as HTML"""
        try:
            from docx import Document
            from pathlib import Path
            import html

            # Get the media record
            media = media_service.get_media_by_id(media_id)
            if not media:
                return '<html><body><h3>Media file not found</h3></body></html>', 404

            # Get file path
            if not media.media_path:
                return '<html><body><h3>File path not available</h3></body></html>', 404

            file_path = Path(media.media_path)
            if not file_path.exists():
                return '<html><body><h3>File not found on disk</h3></body></html>', 404

            # Check file extension
            if file_path.suffix.lower() not in ['.docx', '.doc']:
                return f'<html><body><h3>Unsupported file format: {file_path.suffix}</h3></body></html>', 400

            try:
                # Read DOCX file
                doc = Document(file_path)

                # Convert to HTML
                html_content_parts = []

                # Process each paragraph
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        # Basic styling based on paragraph style
                        style = paragraph.style.name if paragraph.style else 'Normal'
                        if 'Heading' in style:
                            level = style.replace('Heading ', '').strip()
                            if level.isdigit():
                                html_content_parts.append(f'<h{level}>{html.escape(paragraph.text)}</h{level}>')
                            else:
                                html_content_parts.append(f'<h3>{html.escape(paragraph.text)}</h3>')
                        else:
                            html_content_parts.append(f'<p>{html.escape(paragraph.text)}</p>')

                # Process tables
                for table in doc.tables:
                    html_content_parts.append('<table class="table table-bordered table-striped">')
                    for i, row in enumerate(table.rows):
                        html_content_parts.append('<tr>')
                        cell_tag = 'th' if i == 0 else 'td'
                        for cell in row.cells:
                            html_content_parts.append(f'<{cell_tag}>{html.escape(cell.text)}</{cell_tag}>')
                        html_content_parts.append('</tr>')
                    html_content_parts.append('</table>')

                document_html = '\n'.join(html_content_parts)

                # Create complete HTML page
                full_html = f'''
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <title>{media.media_name}</title>
                    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
                    <style>
                        body {{
                            padding: 20px;
                            font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
                            line-height: 1.6;
                        }}
                        .container {{
                            max-width: 900px;
                            margin: 0 auto;
                        }}
                        h1, h2, h3, h4, h5, h6 {{
                            margin-top: 20px;
                            margin-bottom: 10px;
                            color: #212529;
                        }}
                        p {{
                            margin-bottom: 10px;
                        }}
                        table {{
                            margin: 20px 0;
                        }}
                        .header {{
                            border-bottom: 2px solid #dee2e6;
                            padding-bottom: 10px;
                            margin-bottom: 20px;
                        }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <div class="header">
                            <h2>{media.media_name}</h2>
                        </div>
                        <div class="content">
                            {document_html}
                        </div>
                    </div>
                </body>
                </html>
                '''

                return full_html

            except Exception as e:
                return f'<html><body><h3>Error reading file: {str(e)}</h3></body></html>', 500

        except Exception as e:
            return f'<html><body><h3>Error: {str(e)}</h3></body></html>', 500

    @app.route('/media/view/3d/<int:media_id>')
    @login_required
    def view_3d_model(media_id):
        """View 3D models using Three.js"""
        try:
            from pathlib import Path

            # Get the media record
            media = media_service.get_media_by_id(media_id)
            if not media:
                return '<html><body><h3>Media file not found</h3></body></html>', 404

            # Get file path and validate
            if not media.media_path:
                return '<html><body><h3>File path not available</h3></body></html>', 404

            file_path = Path(media.media_path)
            if not file_path.exists():
                return '<html><body><h3>File not found on disk</h3></body></html>', 404

            # Check file extension
            file_ext = file_path.suffix.lower()
            supported_formats = ['.obj', '.stl', '.ply', '.gltf', '.glb', '.dae']
            if file_ext not in supported_formats:
                return f'<html><body><h3>Unsupported 3D format: {file_ext}</h3><p>Supported formats: {", ".join(supported_formats)}</p></body></html>', 400

            # Determine the appropriate Three.js loader based on file extension
            loader_map = {
                '.obj': 'OBJLoader',
                '.stl': 'STLLoader',
                '.ply': 'PLYLoader',
                '.gltf': 'GLTFLoader',
                '.glb': 'GLTFLoader',
                '.dae': 'ColladaLoader'
            }

            loader_type = loader_map.get(file_ext, 'GLTFLoader')

            # Create the model URL - need to serve it relative to web root
            model_url = f'/{media.media_path}'

            # Create complete HTML page with Three.js viewer
            full_html = f'''
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{media.media_name} - 3D Viewer</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>
                    body {{
                        margin: 0;
                        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        overflow: hidden;
                    }}
                    .viewer-container {{
                        position: relative;
                        width: 100vw;
                        height: 100vh;
                        display: flex;
                        flex-direction: column;
                    }}
                    .header {{
                        background: rgba(255, 255, 255, 0.95);
                        padding: 15px 30px;
                        box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
                        z-index: 100;
                    }}
                    .header h2 {{
                        margin: 0;
                        font-size: 1.5rem;
                        color: #333;
                    }}
                    .header .info {{
                        color: #666;
                        font-size: 0.9rem;
                        margin-top: 5px;
                    }}
                    #canvas-container {{
                        flex: 1;
                        position: relative;
                        background: #1a1a2e;
                    }}
                    #loading {{
                        position: absolute;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        text-align: center;
                        color: white;
                        font-size: 1.2rem;
                        z-index: 10;
                    }}
                    .controls {{
                        position: absolute;
                        bottom: 20px;
                        right: 20px;
                        background: rgba(255, 255, 255, 0.9);
                        padding: 15px;
                        border-radius: 10px;
                        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
                        z-index: 10;
                    }}
                    .controls button {{
                        margin: 3px;
                    }}
                    .error {{
                        position: absolute;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        background: rgba(255, 255, 255, 0.95);
                        padding: 30px;
                        border-radius: 10px;
                        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
                        text-align: center;
                        max-width: 500px;
                        z-index: 100;
                    }}
                </style>
            </head>
            <body>
                <div class="viewer-container">
                    <div class="header">
                        <h2><i class="fas fa-cube"></i> {media.media_name}</h2>
                        <div class="info">
                            {media.description if media.description else 'Interactive 3D Model Viewer'}
                        </div>
                    </div>

                    <div id="canvas-container">
                        <div id="loading">
                            <div class="spinner-border text-light" role="status">
                                <span class="visually-hidden">Loading...</span>
                            </div>
                            <p class="mt-3">Loading 3D model...</p>
                        </div>
                    </div>

                    <div class="controls">
                        <div class="mb-2"><small><strong>Controls:</strong></small></div>
                        <div><small>🖱️ Left click + drag: Rotate</small></div>
                        <div><small>🖱️ Right click + drag: Pan</small></div>
                        <div><small>🖱️ Scroll: Zoom</small></div>
                        <div class="mt-2">
                            <button id="resetBtn" class="btn btn-sm btn-primary">
                                <i class="fas fa-redo"></i> Reset View
                            </button>
                            <button id="wireframeBtn" class="btn btn-sm btn-secondary">
                                <i class="fas fa-border-all"></i> Wireframe
                            </button>
                        </div>
                    </div>
                </div>

                <!-- Three.js and loaders (using r147 - last version with legacy format) -->
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/build/three.min.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/controls/OrbitControls.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/loaders/OBJLoader.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/loaders/STLLoader.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/loaders/PLYLoader.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/loaders/GLTFLoader.js"></script>
                <script src="https://cdn.jsdelivr.net/npm/three@0.147.0/examples/js/loaders/ColladaLoader.js"></script>
                <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">

                <script>
                    let scene, camera, renderer, controls, model;
                    let wireframeMode = false;
                    const modelUrl = "{model_url}";
                    const fileExt = "{file_ext}";
                    const loaderType = "{loader_type}";

                    function init() {{
                        const container = document.getElementById('canvas-container');

                        // Scene
                        scene = new THREE.Scene();
                        scene.background = new THREE.Color(0x1a1a2e);

                        // Camera
                        camera = new THREE.PerspectiveCamera(
                            75,
                            container.clientWidth / container.clientHeight,
                            0.1,
                            1000
                        );
                        camera.position.set(5, 5, 5);

                        // Renderer
                        renderer = new THREE.WebGLRenderer({{ antialias: true }});
                        renderer.setSize(container.clientWidth, container.clientHeight);
                        renderer.setPixelRatio(window.devicePixelRatio);
                        container.appendChild(renderer.domElement);

                        // Controls
                        controls = new THREE.OrbitControls(camera, renderer.domElement);
                        controls.enableDamping = true;
                        controls.dampingFactor = 0.05;
                        controls.screenSpacePanning = false;
                        controls.minDistance = 1;
                        controls.maxDistance = 100;

                        // Lights
                        const ambientLight = new THREE.AmbientLight(0xffffff, 0.6);
                        scene.add(ambientLight);

                        const directionalLight1 = new THREE.DirectionalLight(0xffffff, 0.8);
                        directionalLight1.position.set(10, 10, 10);
                        scene.add(directionalLight1);

                        const directionalLight2 = new THREE.DirectionalLight(0xffffff, 0.4);
                        directionalLight2.position.set(-10, -10, -10);
                        scene.add(directionalLight2);

                        // Grid
                        const gridHelper = new THREE.GridHelper(20, 20, 0x444444, 0x222222);
                        scene.add(gridHelper);

                        // Load model
                        loadModel();

                        // Event listeners
                        window.addEventListener('resize', onWindowResize, false);
                        document.getElementById('resetBtn').addEventListener('click', resetCamera);
                        document.getElementById('wireframeBtn').addEventListener('click', toggleWireframe);

                        // Animation loop
                        animate();
                    }}

                    function loadModel() {{
                        const loadingElement = document.getElementById('loading');

                        function onLoad(object) {{
                            // Handle different loader return types
                            if (loaderType === 'GLTFLoader') {{
                                model = object.scene;
                            }} else if (loaderType === 'ColladaLoader') {{
                                model = object.scene;
                            }} else {{
                                // For OBJ, STL, PLY - wrap geometry in mesh
                                const geometry = object.isBufferGeometry ? object : object.geometry;
                                const material = new THREE.MeshPhongMaterial({{
                                    color: 0x00aaff,
                                    shininess: 30,
                                    flatShading: false
                                }});
                                model = new THREE.Mesh(geometry, material);
                            }}

                            // Center and scale model
                            const box = new THREE.Box3().setFromObject(model);
                            const center = box.getCenter(new THREE.Vector3());
                            const size = box.getSize(new THREE.Vector3());
                            const maxDim = Math.max(size.x, size.y, size.z);
                            const scale = 5 / maxDim;

                            model.position.sub(center);
                            model.scale.multiplyScalar(scale);

                            scene.add(model);
                            loadingElement.style.display = 'none';
                        }}

                        function onProgress(xhr) {{
                            const percent = (xhr.loaded / xhr.total) * 100;
                            loadingElement.innerHTML = `
                                <div class="spinner-border text-light" role="status">
                                    <span class="visually-hidden">Loading...</span>
                                </div>
                                <p class="mt-3">Loading 3D model... ${{Math.round(percent)}}%</p>
                            `;
                        }}

                        function onError(error) {{
                            console.error('Error loading model:', error);
                            loadingElement.innerHTML = `
                                <div class="error">
                                    <h3 class="text-danger"><i class="fas fa-exclamation-triangle"></i> Error Loading Model</h3>
                                    <p class="mb-0">Could not load the 3D model file.</p>
                                    <p class="small text-muted mt-2">${{error.message || 'Unknown error'}}</p>
                                </div>
                            `;
                        }}

                        // Choose appropriate loader
                        let loader;
                        switch (loaderType) {{
                            case 'OBJLoader':
                                loader = new THREE.OBJLoader();
                                break;
                            case 'STLLoader':
                                loader = new THREE.STLLoader();
                                break;
                            case 'PLYLoader':
                                loader = new THREE.PLYLoader();
                                break;
                            case 'GLTFLoader':
                                loader = new THREE.GLTFLoader();
                                break;
                            case 'ColladaLoader':
                                loader = new THREE.ColladaLoader();
                                break;
                            default:
                                onError({{ message: 'Unsupported file format' }});
                                return;
                        }}

                        loader.load(modelUrl, onLoad, onProgress, onError);
                    }}

                    function animate() {{
                        requestAnimationFrame(animate);
                        controls.update();
                        renderer.render(scene, camera);
                    }}

                    function onWindowResize() {{
                        const container = document.getElementById('canvas-container');
                        camera.aspect = container.clientWidth / container.clientHeight;
                        camera.updateProjectionMatrix();
                        renderer.setSize(container.clientWidth, container.clientHeight);
                    }}

                    function resetCamera() {{
                        camera.position.set(5, 5, 5);
                        camera.lookAt(0, 0, 0);
                        controls.reset();
                    }}

                    function toggleWireframe() {{
                        wireframeMode = !wireframeMode;
                        if (model) {{
                            model.traverse((child) => {{
                                if (child.isMesh) {{
                                    child.material.wireframe = wireframeMode;
                                }}
                            }});
                        }}
                        const btn = document.getElementById('wireframeBtn');
                        btn.classList.toggle('btn-secondary');
                        btn.classList.toggle('btn-success');
                    }}

                    // Initialize viewer
                    init();
                </script>
            </body>
            </html>
            '''

            return full_html

        except Exception as e:
            return f'<html><body><h3>Error: {{str(e)}}</h3></body></html>', 500

    # Database Administration Routes
    @app.route('/admin/database')
    @login_required
    def admin_database():
        """Database administration page"""
        # Get current database info
        current_url = app.config.get('CURRENT_DATABASE_URL', database_url)

        # Parse current connection info
        db_info = {
            'url': current_url,
            'type': 'SQLite' if current_url.startswith('sqlite') else 'PostgreSQL',
            'is_default': current_url == database_url
        }

        # Get statistics
        try:
            total_sites = site_service.count_sites()
            total_us = us_service.count_us()
            total_inventory = inventario_service.count_inventario()
        except Exception:
            total_sites = total_us = total_inventory = 0

        stats = {
            'sites': total_sites,
            'us': total_us,
            'inventory': total_inventory
        }

        # Load saved connections from ConnectionManager
        from pyarchinit_mini.config.connection_manager import get_connection_manager
        conn_manager = get_connection_manager()
        saved_connections_list = conn_manager.list_connections()

        # Convert to dict format expected by template
        connections = {}
        for conn in saved_connections_list:
            conn_full = conn_manager.get_connection(conn['name'])
            if conn_full:
                connections[conn['name']] = {
                    'type': conn['db_type'],
                    'url': conn_full['connection_string'],
                    'description': conn.get('description', ''),
                    'uploaded': False  # Will be updated if it's an uploaded DB
                }

        return render_template('admin/database.html',
                             db_info=db_info,
                             stats=stats,
                             connections=connections)

    @app.route('/admin/database/upload', methods=['GET', 'POST'])
    def upload_database():
        """Upload SQLite database file"""
        form = DatabaseUploadForm()

        if form.validate_on_submit():
            try:
                uploaded_file = form.database_file.data
                db_name = form.database_name.data
                description = form.description.data

                # Secure filename
                filename = secure_filename(uploaded_file.filename)
                if not filename.endswith('.db'):
                    filename += '.db'

                # Save to databases folder
                db_path = os.path.join(app.config['DATABASE_FOLDER'], filename)
                uploaded_file.save(db_path)

                # Validate it's a valid SQLite database
                try:
                    test_conn = DatabaseConnection.from_url(f'sqlite:///{db_path}')
                    # Try to query tables to validate
                    with test_conn.get_session() as session:
                        session.execute(text('SELECT name FROM sqlite_master WHERE type="table"'))
                except Exception as e:
                    os.remove(db_path)
                    flash(f'File non valido: {str(e)}', 'error')
                    return render_template('admin/database_upload.html', form=form)

                # Store connection info
                connections = app.config.get('DATABASE_CONNECTIONS', {})
                connections[db_name] = {
                    'type': 'sqlite',
                    'path': db_path,
                    'url': f'sqlite:///{db_path}',
                    'description': description,
                    'uploaded': True
                }
                app.config['DATABASE_CONNECTIONS'] = connections

                # Save to persistent ConnectionManager
                from pyarchinit_mini.config.connection_manager import get_connection_manager
                conn_manager = get_connection_manager()
                conn_manager.add_connection(
                    name=db_name,
                    db_type='sqlite',
                    connection_string=f'sqlite:///{db_path}',
                    description=description
                )

                flash(f'Database "{db_name}" caricato con successo!', 'success')
                return redirect(url_for('admin_database'))

            except Exception as e:
                flash(f'Errore caricamento database: {str(e)}', 'error')

        return render_template('admin/database_upload.html', form=form)

    @app.route('/admin/database/connect', methods=['GET', 'POST'])
    def connect_database():
        """Connect to external database (PostgreSQL or local SQLite)"""
        form = DatabaseConnectionForm()

        if form.validate_on_submit():
            try:
                conn_name = form.connection_name.data
                db_type = form.db_type.data

                # Build connection URL
                if db_type == 'postgresql':
                    host = form.host.data or 'localhost'
                    port = form.port.data or 5432
                    database = form.database.data
                    username = form.username.data
                    password = form.password.data

                    if username and password:
                        connection_url = f'postgresql://{username}:{password}@{host}:{port}/{database}'
                    else:
                        connection_url = f'postgresql://{host}:{port}/{database}'

                else:  # SQLite
                    sqlite_path = form.sqlite_path.data
                    if not os.path.exists(sqlite_path):
                        flash(f'File non trovato: {sqlite_path}', 'error')
                        return render_template('admin/database_connect.html', form=form)
                    connection_url = f'sqlite:///{sqlite_path}'

                # Test connection
                try:
                    test_conn = DatabaseConnection.from_url(connection_url)
                    with test_conn.get_session() as session:
                        # Try a simple query
                        session.execute(text('SELECT 1'))
                except Exception as e:
                    flash(f'Errore connessione: {str(e)}', 'error')
                    return render_template('admin/database_connect.html', form=form)

                # Store connection
                connections = app.config.get('DATABASE_CONNECTIONS', {})
                connections[conn_name] = {
                    'type': db_type,
                    'url': connection_url,
                    'uploaded': False
                }
                app.config['DATABASE_CONNECTIONS'] = connections

                # Save to persistent ConnectionManager
                from pyarchinit_mini.config.connection_manager import get_connection_manager
                conn_manager = get_connection_manager()
                conn_manager.add_connection(
                    name=conn_name,
                    db_type=db_type,
                    connection_string=connection_url,
                    description=''
                )

                flash(f'Connessione "{conn_name}" aggiunta con successo!', 'success')
                return redirect(url_for('admin_database'))

            except Exception as e:
                flash(f'Errore: {str(e)}', 'error')

        return render_template('admin/database_connect.html', form=form)

    @app.route('/admin/database/info')
    def database_info():
        """Get detailed information about current database"""
        try:
            # Get table list
            with db_conn.get_session() as session:
                if app.config['CURRENT_DATABASE_URL'].startswith('sqlite'):
                    result = session.execute(text('SELECT name FROM sqlite_master WHERE type="table" ORDER BY name'))
                    tables = [row[0] for row in result]
                else:
                    result = session.execute(text("SELECT tablename FROM pg_tables WHERE schemaname = 'public' ORDER BY tablename"))
                    tables = [row[0] for row in result]

            # Get row counts for main tables
            table_counts = {}
            for table in ['site_table', 'us_table', 'inventario_materiali_table']:
                try:
                    with db_conn.get_session() as session:
                        result = session.execute(text(f'SELECT COUNT(*) FROM {table}'))
                        table_counts[table] = result.scalar()
                except Exception:
                    table_counts[table] = 0

            return render_template('admin/database_info.html',
                                 tables=tables,
                                 table_counts=table_counts,
                                 current_url=app.config['CURRENT_DATABASE_URL'])

        except Exception as e:
            flash(f'Errore recupero info database: {str(e)}', 'error')
            return redirect(url_for('admin_database'))

    @app.route('/admin/database/switch/<connection_name>', methods=['POST'])
    @login_required
    def switch_database(connection_name):
        """Switch active database to a saved connection"""
        try:
            connections = app.config.get('DATABASE_CONNECTIONS', {})

            if connection_name not in connections:
                flash(f'Connessione "{connection_name}" non trovata', 'error')
                return redirect(url_for('admin_database'))

            conn_info = connections[connection_name]
            new_db_url = conn_info['url']

            # Test connection before switching
            try:
                test_conn = DatabaseConnection.from_url(new_db_url)
                with test_conn.get_session() as session:
                    session.execute(text('SELECT 1'))
            except Exception as e:
                flash(f'Errore connessione al database: {str(e)}', 'error')
                return redirect(url_for('admin_database'))

            # Switch database
            app.config['CURRENT_DATABASE_URL'] = new_db_url

            # Close old database connection before switching
            global db_conn, db_manager
            if db_conn:
                db_conn.close()

            # Reinitialize database connection
            db_conn = DatabaseConnection.from_url(new_db_url)
            db_manager = DatabaseManager(db_conn)

            # Reinitialize ALL services with new database manager
            global site_service, us_service, inventario_service, thesaurus_service
            global user_service, analytics_service, relationship_sync_service, datazione_service
            global matrix_generator, export_import_service

            site_service = SiteService(db_manager)
            us_service = USService(db_manager)
            inventario_service = InventarioService(db_manager)
            thesaurus_service = ThesaurusService(db_manager)
            user_service = UserService(db_manager)
            analytics_service = AnalyticsService(db_manager)
            relationship_sync_service = RelationshipSyncService(db_manager)
            datazione_service = DatazioneService(db_manager)
            matrix_generator = HarrisMatrixGenerator(db_manager, us_service)
            export_import_service = ImportExportService(db_manager.connection.connection_string)

            # Update app stored services (ALL of them to ensure switch works)
            app.db_manager = db_manager
            app.user_service = user_service
            app.site_service = site_service
            app.us_service = us_service
            app.inventario_service = inventario_service
            app.thesaurus_service = thesaurus_service
            app.analytics_service = analytics_service
            app.relationship_sync_service = relationship_sync_service
            app.datazione_service = datazione_service
            app.matrix_generator = matrix_generator
            app.export_import_service = export_import_service
            app.media_service = media_service

            # Log the switch for debugging
            print(f"[DATABASE SWITCH] Changed to: {new_db_url}")
            print(f"[DATABASE SWITCH] Connection name: {connection_name}")

            flash(f'Database cambiato a: {connection_name}', 'success')

        except Exception as e:
            flash(f'Errore cambio database: {str(e)}', 'error')

        return redirect(url_for('admin_database'))

    @app.route('/admin/database/remove/<connection_name>', methods=['POST'])
    @login_required
    def remove_connection(connection_name):
        """Remove a saved connection"""
        try:
            # Check if trying to remove active connection
            current_url = app.config.get('CURRENT_DATABASE_URL')
            connections = app.config.get('DATABASE_CONNECTIONS', {})

            if connection_name in connections:
                if connections[connection_name]['url'] == current_url:
                    flash('Non puoi eliminare la connessione attualmente in uso', 'error')
                    return redirect(url_for('admin_database'))

            # Remove from ConnectionManager
            from pyarchinit_mini.config.connection_manager import get_connection_manager
            conn_manager = get_connection_manager()
            result = conn_manager.remove_connection(connection_name)

            if result['success']:
                # Also remove from app.config
                if connection_name in connections:
                    del connections[connection_name]
                    app.config['DATABASE_CONNECTIONS'] = connections

                flash(f'Connessione "{connection_name}" eliminata con successo', 'success')
            else:
                flash(f'Errore: {result["message"]}', 'error')

        except Exception as e:
            flash(f'Errore eliminazione connessione: {str(e)}', 'error')

        return redirect(url_for('admin_database'))

    # API endpoints for AJAX
    @app.route('/api/sites')
    def api_sites():
        sites = site_service.get_all_sites(size=100)
        return jsonify([{'id': s.id_sito, 'name': s.sito} for s in sites])

    # ===== Export/Import Routes =====

    @app.route('/export')
    @login_required
    def export_page():
        """Export/Import management page"""
        return render_template('export/export_import.html')

    # Excel Export Routes
    @app.route('/export/sites/excel')
    @login_required
    def export_sites_excel():
        """Export sites to Excel"""
        try:
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
            export_import_service.export_sites_to_excel(output_path)
            return send_file(output_path, as_attachment=True,
                           download_name='sites_export.xlsx',
                           mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        except Exception as e:
            flash(f'Errore export Excel: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/export/us/excel')
    def export_us_excel():
        """Export US to Excel"""
        try:
            site_name = request.args.get('sito')
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
            export_import_service.export_us_to_excel(output_path, site_name)
            filename = f'us_{site_name}.xlsx' if site_name else 'us_export.xlsx'
            return send_file(output_path, as_attachment=True,
                           download_name=filename,
                           mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        except Exception as e:
            flash(f'Errore export Excel: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/export/inventario/excel')
    def export_inventario_excel():
        """Export Inventario to Excel"""
        try:
            site_name = request.args.get('sito')
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
            export_import_service.export_inventario_to_excel(output_path, site_name)
            filename = f'inventario_{site_name}.xlsx' if site_name else 'inventario_export.xlsx'
            return send_file(output_path, as_attachment=True,
                           download_name=filename,
                           mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        except Exception as e:
            flash(f'Errore export Excel: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    # CSV Export Routes
    @app.route('/export/sites/csv')
    def export_sites_csv():
        """Export sites to CSV"""
        try:
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.csv').name
            export_import_service.export_sites_to_csv(output_path)
            return send_file(output_path, as_attachment=True,
                           download_name='sites_export.csv',
                           mimetype='text/csv')
        except Exception as e:
            flash(f'Errore export CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/export/us/csv')
    def export_us_csv():
        """Export US to CSV"""
        try:
            site_name = request.args.get('sito')
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.csv').name
            export_import_service.export_us_to_csv(output_path, site_name)
            filename = f'us_{site_name}.csv' if site_name else 'us_export.csv'
            return send_file(output_path, as_attachment=True,
                           download_name=filename,
                           mimetype='text/csv')
        except Exception as e:
            flash(f'Errore export CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/export/inventario/csv')
    def export_inventario_csv():
        """Export Inventario to CSV"""
        try:
            site_name = request.args.get('sito')
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.csv').name
            export_import_service.export_inventario_to_csv(output_path, site_name)
            filename = f'inventario_{site_name}.csv' if site_name else 'inventario_export.csv'
            return send_file(output_path, as_attachment=True,
                           download_name=filename,
                           mimetype='text/csv')
        except Exception as e:
            flash(f'Errore export CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    # CSV Import Routes
    @app.route('/import/sites/csv', methods=['POST'])
    def import_sites_csv():
        """Import sites from CSV"""
        try:
            if 'file' not in request.files:
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            file = request.files['file']
            if file.filename == '':
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            if not file.filename.endswith('.csv'):
                flash('Il file deve essere in formato CSV', 'error')
                return redirect(url_for('export_page'))

            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name

            # Import data
            skip_duplicates = request.form.get('skip_duplicates', 'true').lower() == 'true'
            result = export_import_service.batch_import_sites_from_csv(tmp_path, skip_duplicates)

            # Clean up
            os.unlink(tmp_path)

            flash(f'Import completato: {result["imported"]} importati, '
                  f'{result["skipped"]} saltati, {len(result["errors"])} errori', 'success')

            if result['errors']:
                for err in result['errors'][:5]:  # Show first 5 errors
                    flash(f'Errore: {err["error"]}', 'warning')

            return redirect(url_for('export_page'))

        except Exception as e:
            flash(f'Errore import CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/import/us/csv', methods=['POST'])
    def import_us_csv():
        """Import US from CSV"""
        try:
            if 'file' not in request.files:
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            file = request.files['file']
            if file.filename == '':
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            if not file.filename.endswith('.csv'):
                flash('Il file deve essere in formato CSV', 'error')
                return redirect(url_for('export_page'))

            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name

            # Import data
            skip_duplicates = request.form.get('skip_duplicates', 'true').lower() == 'true'
            result = export_import_service.batch_import_us_from_csv(tmp_path, skip_duplicates)

            # Clean up
            os.unlink(tmp_path)

            flash(f'Import completato: {result["imported"]} importati, '
                  f'{result["skipped"]} saltati, {len(result["errors"])} errori', 'success')

            if result['errors']:
                for err in result['errors'][:5]:
                    flash(f'Errore: {err["error"]}', 'warning')

            return redirect(url_for('export_page'))

        except Exception as e:
            flash(f'Errore import CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    @app.route('/import/inventario/csv', methods=['POST'])
    def import_inventario_csv():
        """Import Inventario from CSV"""
        try:
            if 'file' not in request.files:
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            file = request.files['file']
            if file.filename == '':
                flash('Nessun file selezionato', 'error')
                return redirect(url_for('export_page'))

            if not file.filename.endswith('.csv'):
                flash('Il file deve essere in formato CSV', 'error')
                return redirect(url_for('export_page'))

            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name

            # Import data
            skip_duplicates = request.form.get('skip_duplicates', 'true').lower() == 'true'
            result = export_import_service.batch_import_inventario_from_csv(tmp_path, skip_duplicates)

            # Clean up
            os.unlink(tmp_path)

            flash(f'Import completato: {result["imported"]} importati, '
                  f'{result["skipped"]} saltati, {len(result["errors"])} errori', 'success')

            if result['errors']:
                for err in result['errors'][:5]:
                    flash(f'Errore: {err["error"]}', 'warning')

            return redirect(url_for('export_page'))

        except Exception as e:
            flash(f'Errore import CSV: {str(e)}', 'error')
            return redirect(url_for('export_page'))

    # ===== Periodizzazione (Datazioni) Routes =====
    @app.route('/periodizzazione')
    @login_required
    def periodizzazione_list():
        """List all datazioni"""
        try:
            datazioni = datazione_service.get_all_datazioni()
            total = datazione_service.count_datazioni()

            return render_template('datazioni/list.html',
                                 datazioni=datazioni,
                                 total=total)
        except Exception as e:
            flash(f'Errore caricamento datazioni: {str(e)}', 'error')
            return render_template('datazioni/list.html',
                                 datazioni=[],
                                 total=0)

    @app.route('/periodizzazione/create', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def periodizzazione_create():
        """Create new datazione"""
        if request.method == 'POST':
            try:
                datazione_data = {
                    'nome_datazione': request.form.get('nome_datazione'),
                    'fascia_cronologica': request.form.get('fascia_cronologica'),
                    'descrizione': request.form.get('descrizione')
                }

                datazione_service.create_datazione(datazione_data)
                flash('Datazione creata con successo!', 'success')
                return redirect(url_for('periodizzazione_list'))

            except Exception as e:
                flash(f'Errore creazione datazione: {str(e)}', 'error')

        return render_template('periodizzazione/form.html',
                             datazione=None,
                             title='New Dating')

    @app.route('/periodizzazione/<int:datazione_id>/edit', methods=['GET', 'POST'])
    @login_required
    @write_permission_required
    def periodizzazione_edit(datazione_id):
        """Edit existing datazione"""
        try:
            datazione = datazione_service.get_datazione_by_id(datazione_id)
            if not datazione:
                flash('Datazione non trovata', 'error')
                return redirect(url_for('periodizzazione_list'))

            if request.method == 'POST':
                update_data = {
                    'nome_datazione': request.form.get('nome_datazione'),
                    'fascia_cronologica': request.form.get('fascia_cronologica'),
                    'descrizione': request.form.get('descrizione')
                }

                datazione_service.update_datazione(datazione_id, update_data)
                flash('Datazione aggiornata con successo!', 'success')
                return redirect(url_for('periodizzazione_list'))

            return render_template('periodizzazione/form.html',
                                 datazione=datazione,
                                 title='Edit Dating')

        except Exception as e:
            flash(f'Errore modifica datazione: {str(e)}', 'error')
            return redirect(url_for('periodizzazione_list'))

    @app.route('/periodizzazione/<int:datazione_id>/delete', methods=['POST'])
    @login_required
    @write_permission_required
    def periodizzazione_delete(datazione_id):
        """Delete datazione"""
        try:
            datazione_service.delete_datazione(datazione_id)
            flash('Datazione eliminata con successo!', 'success')
        except Exception as e:
            flash(f'Errore eliminazione datazione: {str(e)}', 'error')

        return redirect(url_for('periodizzazione_list'))

    # ===== Periodization Records Routes =====
    @app.route('/periodization-records')
    @login_required
    def periodization_records_list():
        """List all periodization records from periodizzazione_table"""
        try:
            from pyarchinit_mini.models.harris_matrix import Periodizzazione
            from sqlalchemy import or_

            # Get search parameters
            search_site = request.args.get('search_site', '')
            search_us = request.args.get('search_us', '')
            search_period = request.args.get('search_period', '')
            page = request.args.get('page', 1, type=int)
            page_size = 50

            with db_manager.connection.get_session() as session:
                query = session.query(Periodizzazione)

                # Apply filters
                if search_site:
                    query = query.filter(Periodizzazione.sito.contains(search_site))
                if search_us:
                    query = query.filter(Periodizzazione.us == int(search_us)) if search_us.isdigit() else query
                if search_period:
                    query = query.filter(or_(
                        Periodizzazione.periodo_iniziale.contains(search_period),
                        Periodizzazione.periodo_finale.contains(search_period),
                        Periodizzazione.datazione_estesa.contains(search_period)
                    ))

                # Get total count
                total = query.count()

                # Apply pagination and ordering
                db_records = query.order_by(Periodizzazione.sito, Periodizzazione.us)\
                              .offset((page - 1) * page_size)\
                              .limit(page_size)\
                              .all()

                # Convert to dicts to avoid detached instance errors
                records = [{
                    'sito': r.sito,
                    'area': r.area,
                    'us': r.us,
                    'periodo_iniziale': r.periodo_iniziale,
                    'fase_iniziale': r.fase_iniziale,
                    'periodo_finale': r.periodo_finale,
                    'fase_finale': r.fase_finale,
                    'datazione_estesa': r.datazione_estesa,
                    'affidabilita': r.affidabilita
                } for r in db_records]

            return render_template('periodizzazione/periods.html',
                                 records=records,
                                 total=total,
                                 page=page,
                                 page_size=page_size,
                                 search_site=search_site,
                                 search_us=search_us,
                                 search_period=search_period)

        except Exception as e:
            flash(f'Error loading periodization records: {str(e)}', 'error')
            return render_template('periodizzazione/periods.html',
                                 records=[],
                                 total=0,
                                 page=1,
                                 page_size=50)

    # ===== Thesaurus Routes =====
    @app.route('/thesaurus')
    @login_required
    def thesaurus_list():
        """List thesaurus by table and field"""
        try:
            # Get available tables from THESAURUS_MAPPINGS
            from pyarchinit_mini.models.thesaurus import THESAURUS_MAPPINGS

            tables = list(THESAURUS_MAPPINGS.keys())

            # Get selected table and field from query params
            selected_table = request.args.get('table', '')
            selected_field = request.args.get('field', '')

            # Get fields for selected table
            fields = []
            values = []
            if selected_table:
                fields = thesaurus_service.get_table_fields(selected_table)

                # Get values for selected field
                if selected_field:
                    values = thesaurus_service.get_field_values(selected_table, selected_field)

            return render_template('thesaurus/list.html',
                                 tables=tables,
                                 fields=fields,
                                 values=values,
                                 selected_table=selected_table,
                                 selected_field=selected_field)

        except Exception as e:
            flash(f'Thesaurus loading error: {str(e)}', 'error')
            return render_template('thesaurus/list.html',
                                 tables=[],
                                 fields=[],
                                 values=[],
                                 selected_table='',
                                 selected_field='')

    @app.route('/thesaurus/create', methods=['POST'])
    @login_required
    @write_permission_required
    def thesaurus_create():
        """Create new thesaurus value"""
        try:
            table_name = request.form.get('table_name')
            field_name = request.form.get('field_name')
            value = request.form.get('value')
            label = request.form.get('label')
            description = request.form.get('description')

            thesaurus_service.add_field_value(
                table_name=table_name,
                field_name=field_name,
                value=value,
                label=label,
                description=description
            )

            flash('Thesaurus value created successfully!', 'success')

        except Exception as e:
            flash(f'Value creation error: {str(e)}', 'error')

        return redirect(url_for('thesaurus_list',
                               table=request.form.get('table_name'),
                               field=request.form.get('field_name')))

    @app.route('/thesaurus/<field_id>/edit', methods=['POST'])
    @login_required
    @write_permission_required
    def thesaurus_edit(field_id):
        """Edit thesaurus value"""
        try:
            # Check if this is a predefined value (read-only)
            if str(field_id).startswith('predefined_'):
                flash('Predefined values cannot be modified. Create a new custom value.', 'warning')
            else:
                value = request.form.get('value')
                label = request.form.get('label')
                description = request.form.get('description')

                thesaurus_service.update_field_value(
                    field_id=int(field_id),
                    value=value,
                    label=label,
                    description=description
                )

                flash('Thesaurus value updated successfully!', 'success')

        except Exception as e:
            flash(f'Value update error: {str(e)}', 'error')

        table_name = request.form.get('table_name')
        field_name = request.form.get('field_name')
        return redirect(url_for('thesaurus_list',
                               table=table_name,
                               field=field_name))

    @app.route('/thesaurus/<field_id>/delete', methods=['POST'])
    @login_required
    @write_permission_required
    def thesaurus_delete(field_id):
        """Delete thesaurus value"""
        try:
            # Check if this is a predefined value (read-only)
            if str(field_id).startswith('predefined_'):
                flash('Predefined values cannot be deleted. They are read-only.', 'warning')
            else:
                thesaurus_service.delete_field_value(int(field_id))
                flash('Thesaurus value deleted successfully!', 'success')
        except Exception as e:
            flash(f'Value deletion error: {str(e)}', 'error')

        table_name = request.form.get('table_name')
        field_name = request.form.get('field_name')
        return redirect(url_for('thesaurus_list',
                               table=table_name,
                               field=field_name))

    # ===== 3D Model Viewer Routes (s3Dgraphy Integration) =====
    try:
        from s3d_routes import init_s3d_routes
        init_s3d_routes(app, db_manager, media_handler)
        print("[FLASK] s3Dgraphy routes initialized")
    except ImportError as e:
        print(f"[FLASK] Warning: s3Dgraphy routes not available: {e}")
    except Exception as e:
        print(f"[FLASK] Error initializing s3Dgraphy routes: {e}")

    return app, socketio

# Run app
def main():
    """
    Entry point for running the web interface via console script.
    """
    app, socketio = create_app()

    # Get configuration from environment or use defaults
    host = os.getenv("PYARCHINIT_WEB_HOST", "0.0.0.0")
    port = int(os.getenv("PYARCHINIT_WEB_PORT", "5001"))  # Changed from 5000 to 5001 to avoid macOS AirPlay
    debug = os.getenv("PYARCHINIT_WEB_DEBUG", "true").lower() == "true"

    print(f"Starting PyArchInit-Mini Web Interface on {host}:{port}")
    print(f"Web Interface: http://{host if host != '0.0.0.0' else 'localhost'}:{port}/")
    print(f"WebSocket support enabled for real-time collaboration")

    # Note: For production use, deploy with gunicorn + eventlet/gevent
    # Example: gunicorn --worker-class eventlet -w 1 --bind 0.0.0.0:5001 app:app
    socketio.run(app, debug=debug, host=host, port=port, allow_unsafe_werkzeug=True)


if __name__ == '__main__':
    main()